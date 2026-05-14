#!/usr/bin/env python3
"""
SDLC Control Validator - Template JSON + Template PDF Evidence Validation
Validates solution design documentation using text embeddings and LLM-based quality assessment.
"""

import os
import io
import json
import ast
import argparse
import logging
import hashlib
import base64
import atexit
import re
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from xml.sax.saxutils import escape as xml_escape

# Database & Connectors
import sqlalchemy
from sqlalchemy import text
from google.cloud.sql.connector import Connector, IPTypes

# Vertex AI (text embed LLM)
from google import genai
from google.genai import types as genai_types
import vertexai

# Auth for REST predict (multimodal embed)
import google.auth
from google.auth.transport.requests import Request as GARequest
import requests

# PDF/Docx parsing (Template PDF guidance text)
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# Local text chunking & cosine similarity for picking best template-PDF snippets
import numpy as np
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sklearn.metrics.pairwise import cosine_similarity

# PDF Generator
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, Spacer


# ============================================================================
# CONFIG / CONSTANTS
# ============================================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s"
)

# Vertex AI Configuration
PROJECT_ID = os.getenv("VERTEX_PROJECT_ID")
REGION = os.getenv("VERTEX_REGION", "europe-west3")

# Embedding & LLM Models
EMBED_MODEL = os.getenv("VERTEX_EMBED_MODEL", "text-embedding-005")
LLM_MODEL = os.getenv("VERTEX_LLM_MODEL", "gemini-2.5-flash")
MM_EMBED_MODEL = os.getenv("VERTEX_MM_EMBED_MODEL", "multimodalembedding@001")

# Embedding Dimensions
MM_DIM = int(os.getenv("VERTEX_MM_DIM", "1408"))

# Cloud SQL Connector Configuration
INSTANCE_CONNECTION_NAME = os.getenv("INSTANCE_CONNECTION_NAME", "")
PGUSER = os.getenv("PGUSER", os.getenv("PG_USER", "master"))
PGPASSWORD = os.getenv("PGPASSWORD", os.getenv("PG_PASSWORD", ""))
PGDATABASE = os.getenv("PGDATABASE", os.getenv("PG_DB", "database"))

# Database Table Names (aligned with your loaders & DDL)
TABLE_EVIDENCE = "vs_evidence"  # text evidence chunks
TABLE_GUIDANCE = "vs_guidance"  # guidance text chunks
TABLE_MM_EVID = "vs_mm_evidence_assets"  # evidence images

# Embedding Dimensions Map
EMBED_DIM_MAP: Dict[str, int] = {
    "text-embedding-005": 768,
    "gemini-embedding-001": 3072,
}
EMBED_DIM: Optional[int] = EMBED_DIM_MAP.get(EMBED_MODEL)

# Allowed Release Types
ALLOWED_RTYPES = {
    "indev": "New (indev) releases",
    "new": "GCP New releases",
    "migration": "GCP Migration releases",
    "rehost": "GCP Re-Host releases",
    "normal": "GCP Normal releases",
}


def resolve_rtype(user_val: str) -> str:
    """Resolve user-provided rtype to canonical form."""
    key = (user_val or "").strip().lower()
    if key in ALLOWED_RTYPES:
        return ALLOWED_RTYPES[key]
    for canonical in ALLOWED_RTYPES.values():
        if key == canonical.lower():
            return canonical
    raise ValueError("Invalid --rtype. Allowed: indev, new, migration, rehost, normal")


# ============================================================================
# Vertex AI Initialization
# ============================================================================

if not PROJECT_ID:
    logging.warning("VERTEX_PROJECT_ID is not set.")

vertexai.init(project=PROJECT_ID, location=REGION)
_genai = genai.Client(vertexai=True, project=PROJECT_ID, location=REGION)


# ============================================================================
# Database: Cloud SQL Connector
# ============================================================================

_connector: Optional[Connector] = None
_engine: Optional[sqlalchemy.Engine] = None


def build_engine() -> sqlalchemy.Engine:
    """Build SQLAlchemy engine via Cloud SQL Connector."""
    global _connector
    _connector = Connector()

    def getconn():
        return _connector.connect(
            INSTANCE_CONNECTION_NAME,
            "pg8000",
            user=PGUSER,
            password=PGPASSWORD,
            db=PGDATABASE,
            ip_type=IPTypes.PSC,
        )

    return sqlalchemy.create_engine(
        "postgresql+pg8000://",
        creator=getconn,
        pool_pre_ping=True,
        pool_recycle=300,
    )


def get_engine() -> sqlalchemy.Engine:
    """Get or create the shared SQLAlchemy engine."""
    global _engine
    if _engine is None:
        if not INSTANCE_CONNECTION_NAME:
            raise RuntimeError("INSTANCE_CONNECTION_NAME is required for Cloud SQL Connector.")
        _engine = build_engine()
        logging.info("SQLAlchemy engine created via Cloud SQL Connector (pg8000).")
    return _engine


def _get_column_format_type(table_name: str, column_name: str) -> Optional[str]:
    """Return PostgreSQL's formatted column type, e.g. vector(1408)."""
    engine = get_engine()
    sql = text("""
        SELECT format_type(a.atttypid, a.atttypmod) AS format_type
        FROM pg_attribute a
        JOIN pg_class c ON c.oid = a.attrelid
        JOIN pg_namespace n ON n.oid = c.relnamespace
        WHERE c.relname = :table_name
          AND a.attname = :column_name
          AND a.attnum > 0
          AND NOT a.attisdropped
        ORDER BY CASE WHEN n.nspname = current_schema() THEN 0 ELSE 1 END, n.nspname
        LIMIT 1
    """)
    with engine.connect() as conn:
        row = conn.execute(sql, {"table_name": table_name, "column_name": column_name}).first()
    return row[0] if row and row[0] else None


def _parse_vector_dimension(format_type_name: Optional[str]) -> Optional[int]:
    """Parse a vector(N) type string into its dimension."""
    if not format_type_name:
        return None
    match = re.fullmatch(r"vector\((\d+)\)", format_type_name.strip(), flags=re.IGNORECASE)
    return int(match.group(1)) if match else None


def validate_runtime_contract(mm_dim: int) -> None:
    """Fail fast when runtime configuration does not match the database contract."""
    if EMBED_DIM is not None:
        text_embedding_type = _get_column_format_type(TABLE_EVIDENCE, "embedding")
        text_embedding_dim = _parse_vector_dimension(text_embedding_type)
        if text_embedding_dim != EMBED_DIM:
            raise RuntimeError(
                f"{TABLE_EVIDENCE}.embedding expects {text_embedding_type or 'unknown'}, "
                f"but {EMBED_MODEL} is configured for dimension {EMBED_DIM}."
            )

    guidance_embedding_type = _get_column_format_type(TABLE_GUIDANCE, "embedding")
    guidance_embedding_dim = _parse_vector_dimension(guidance_embedding_type)
    if guidance_embedding_dim is not None and EMBED_DIM is not None and guidance_embedding_dim != EMBED_DIM:
        raise RuntimeError(
            f"{TABLE_GUIDANCE}.embedding expects {guidance_embedding_type}, "
            f"but {EMBED_MODEL} is configured for dimension {EMBED_DIM}."
        )

    if not mm_dim:
        return

    mm_embedding_type = _get_column_format_type(TABLE_MM_EVID, "embedding")
    mm_embedding_dim = _parse_vector_dimension(mm_embedding_type)
    if mm_embedding_dim != mm_dim:
        raise RuntimeError(
            f"{TABLE_MM_EVID}.embedding expects {mm_embedding_type or 'unknown'}, "
            f"but VERTEX_MM_DIM is configured for dimension {mm_dim}."
        )


def shutdown_connector():
    """Cleanup: close Cloud SQL connector."""
    global _connector, _engine
    try:
        if _engine is not None:
            _engine.dispose()
    finally:
        if _connector is not None:
            _connector.close()
        logging.info("Cloud SQL Connector closed.")


atexit.register(shutdown_connector)


# ============================================================================
# Embedding & LLM Helpers
# ============================================================================

def _extract_values(resp) -> List[float]:
    """Extract embedding vector from various response formats."""
    if hasattr(resp, "embedding") and hasattr(resp.embedding, "values"):
        return resp.embedding.values
    if hasattr(resp, "values"):
        return resp.values
    if hasattr(resp, "embeddings") and resp.embeddings:
        for e in resp.embeddings:
            if hasattr(e, "values"):
                return e.values
            if isinstance(e, dict) and "values" in e:
                return e["values"]
    if isinstance(resp, dict):
        if "embedding" in resp and isinstance(resp["embedding"], dict) and "values" in resp["embedding"]:
            return resp["embedding"]["values"]
        if "values" in resp:
            return resp["values"]
    raise RuntimeError(f"Unexpected embed response shape: {type(resp)}")


def embed_texts(
    texts: List[str],
    task_type: str = "RETRIEVAL_DOCUMENT",
    out_dim: Optional[int] = EMBED_DIM,
) -> List[List[float]]:
    """Embed multiple text chunks."""
    cfg = genai_types.EmbedContentConfig(
        task_type=task_type,
        output_dimensionality=out_dim,
    )
    out = []
    for t in texts:
        r = _genai.models.embed_content(model=EMBED_MODEL, contents=t, config=cfg)
        out.append(_extract_values(r))
    return out


def embed_query_text(
    text_q: str,
    out_dim: Optional[int] = EMBED_DIM,
) -> List[float]:
    """Embed a query text."""
    cfg = genai_types.EmbedContentConfig(
        task_type="RETRIEVAL_QUERY",
        output_dimensionality=out_dim,
    )
    r = _genai.models.embed_content(model=EMBED_MODEL, contents=text_q, config=cfg)
    return _extract_values(r)


def get_access_token() -> str:
    """Get Google Cloud access token for REST API calls."""
    creds, _ = google.auth.default(scopes=["https://www.googleapis.com/auth/cloud-platform"])
    if not creds.valid:
        creds.refresh(GARequest())
    return creds.token


def mm_embed_text(query: str, dim: int = MM_DIM) -> List[float]:
    """Embed text using multimodal embedding model (REST API)."""
    token = get_access_token()
    endpoint = (
        f"https://{REGION}-aiplatform.googleapis.com/v1/projects/{PROJECT_ID}/locations/{REGION}"
        f"/publishers/google/models/{MM_EMBED_MODEL}:predict"
    )
    payload = {
        "instances": [
            {
                "text": query,
                "parameters": {"dimension": dim},
            }
        ]
    }
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }
    resp = requests.post(endpoint, headers=headers, json=payload, timeout=60)
    resp.raise_for_status()
    data = resp.json()
    preds = data.get("predictions") or []
    if not preds or "textEmbedding" not in preds[0]:
        raise RuntimeError(f"Vertex MM: expected textEmbedding. data_keys={list(data.keys())}")
    vec = preds[0]["textEmbedding"]
    if dim and len(vec) != dim:
        raise RuntimeError(f"Vertex MM: expected dim={dim}, got len={len(vec)}")
    return vec


def mm_embed_image_with_caption(
    img_bytes: bytes,
    caption: str,
    dim: int = MM_DIM,
) -> List[float]:
    """Embed image with optional caption using multimodal embedding model."""
    token = get_access_token()
    endpoint = (
        f"https://{REGION}-aiplatform.googleapis.com/v1/projects/{PROJECT_ID}/locations/{REGION}"
        f"/publishers/google/models/{MM_EMBED_MODEL}:predict"
    )
    b64 = base64.b64encode(img_bytes).decode("utf-8")
    instance = {
        "image": {
            "bytesBase64Encoded": b64,
            "mimeType": "image/png",
        },
        "parameters": {"dimension": dim},
    }
    if caption:
        instance["text"] = caption[:200]
    payload = {"instances": [instance]}
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }
    resp = requests.post(endpoint, headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    data = resp.json()
    preds = data.get("predictions") or []
    if not preds:
        raise RuntimeError("Vertex MM: missing predictions")
    pred = preds[0]
    vec = pred.get("imageEmbedding") or pred.get("textEmbedding")
    if vec is None:
        raise RuntimeError(f"Vertex MM: missing embedding. keys={list(pred.keys())}")
    if dim and len(vec) != dim:
        raise RuntimeError(f"Vertex MM: expected dim={dim}, got len={len(vec)}")
    return vec


def _safe_getattr(value: Any, attr_name: str, default: Any = None) -> Any:
    """Read SDK attributes without letting property access abort extraction."""
    try:
        return getattr(value, attr_name)
    except Exception:
        return default


def _to_plain_response_data(value: Any) -> Any:
    """Best-effort conversion of SDK objects into plain dict/list data."""
    if value is None or isinstance(value, (str, int, float, bool, dict, list, tuple)):
        return value

    for method_name, kwargs in (
        ("model_dump", {"exclude_none": True}),
        ("to_json_dict", {}),
        ("dict", {}),
    ):
        method = _safe_getattr(value, method_name)
        if not callable(method):
            continue
        try:
            return method(**kwargs)
        except TypeError:
            try:
                return method()
            except Exception:
                continue
        except Exception:
            continue

    try:
        data = {k: v for k, v in vars(value).items() if not k.startswith("_")}
    except Exception:
        data = None
    return data if data else value


def _looks_like_json_payload(text_value: str) -> bool:
    """Identify candidate strings that look like model JSON output."""
    stripped = (text_value or "").strip()
    return (
        stripped.startswith("{")
        or stripped.startswith("[")
        or stripped.startswith("```json")
        or stripped.startswith("```")
    )


def _collect_response_text_candidates(value: Any, depth: int = 0) -> List[str]:
    """Recursively collect text payloads from SDK response objects."""
    if depth > 5 or value is None:
        return []

    value = _to_plain_response_data(value)
    if isinstance(value, str):
        stripped = value.strip()
        return [stripped] if stripped else []

    if isinstance(value, (list, tuple)):
        texts: List[str] = []
        for item in value:
            texts.extend(_collect_response_text_candidates(item, depth + 1))
        return texts

    if isinstance(value, dict):
        for key in ("output_text", "text"):
            text_value = value.get(key)
            if isinstance(text_value, str) and text_value.strip():
                return [text_value.strip()]

        parsed = value.get("parsed")
        if parsed is not None:
            parsed = _to_plain_response_data(parsed)
            if isinstance(parsed, (dict, list)):
                return [json.dumps(parsed, ensure_ascii=False)]
            if isinstance(parsed, str) and parsed.strip():
                return [parsed.strip()]

        for key in ("parts", "content", "candidates"):
            child_texts = _collect_response_text_candidates(value.get(key), depth + 1)
            if child_texts:
                return child_texts

        texts: List[str] = []
        for child in value.values():
            texts.extend(_collect_response_text_candidates(child, depth + 1))
        return texts

    return []


def _summarize_generation_response(resp: Any) -> str:
    """Return compact response metadata for empty-output diagnostics."""
    payload = _to_plain_response_data(resp)
    summary_bits = [f"response_type={type(resp).__name__}"]
    if not isinstance(payload, dict):
        return ", ".join(summary_bits)

    candidates = payload.get("candidates") or []
    summary_bits.append(f"candidate_count={len(candidates)}")

    finish_reasons: List[str] = []
    for candidate in candidates:
        candidate_payload = _to_plain_response_data(candidate)
        if not isinstance(candidate_payload, dict):
            continue
        finish_reason = candidate_payload.get("finish_reason") or candidate_payload.get("finishReason")
        if finish_reason:
            finish_reasons.append(str(finish_reason))
    if finish_reasons:
        summary_bits.append(f"finish_reasons={','.join(sorted(set(finish_reasons)))}")

    prompt_feedback = payload.get("prompt_feedback") or payload.get("promptFeedback")
    prompt_feedback = _to_plain_response_data(prompt_feedback)
    if isinstance(prompt_feedback, dict):
        block_reason = prompt_feedback.get("block_reason") or prompt_feedback.get("blockReason")
        if block_reason:
            summary_bits.append(f"block_reason={block_reason}")

    return ", ".join(summary_bits)


def _extract_generated_text(resp: Any) -> str:
    """Extract generated text from different SDK response shapes."""
    direct_text = _safe_getattr(resp, "output_text") or _safe_getattr(resp, "text")
    if isinstance(direct_text, str) and direct_text.strip():
        return direct_text.strip()

    parsed = _safe_getattr(resp, "parsed")
    parsed = _to_plain_response_data(parsed)
    if isinstance(parsed, (dict, list)):
        return json.dumps(parsed, ensure_ascii=False)
    if isinstance(parsed, str) and parsed.strip():
        return parsed.strip()

    candidate_texts: List[str] = []
    candidates = _safe_getattr(resp, "candidates")
    if candidates:
        for candidate in candidates:
            content = _safe_getattr(candidate, "content")
            parts = _safe_getattr(content, "parts") if content is not None else None
            if not parts:
                candidate_texts.extend(
                    text
                    for text in _collect_response_text_candidates(candidate)
                    if _looks_like_json_payload(text)
                )
                continue
            for part in parts:
                part_text = _safe_getattr(part, "text")
                if isinstance(part_text, str) and part_text.strip():
                    candidate_texts.append(part_text.strip())
                    continue
                candidate_texts.extend(
                    text
                    for text in _collect_response_text_candidates(part)
                    if _looks_like_json_payload(text)
                )
        if candidate_texts:
            return "\n".join(dict.fromkeys(candidate_texts))

    if isinstance(resp, dict):
        if isinstance(resp.get("text"), str) and resp["text"].strip():
            return resp["text"].strip()
        for candidate in resp.get("candidates") or []:
            content = candidate.get("content") if isinstance(candidate, dict) else None
            parts = content.get("parts") if isinstance(content, dict) else None
            if not isinstance(parts, list):
                continue
            for part in parts:
                if isinstance(part, dict) and isinstance(part.get("text"), str) and part["text"].strip():
                    candidate_texts.append(part["text"].strip())
        if candidate_texts:
            return "\n".join(candidate_texts)

    fallback_texts = [
        text
        for text in _collect_response_text_candidates(resp)
        if _looks_like_json_payload(text)
    ]
    if fallback_texts:
        return "\n".join(dict.fromkeys(fallback_texts))

    return ""


def _generate_with_config(prompt: str, generation_config: genai_types.GenerateContentConfig) -> Tuple[str, Any]:
    """Run a single generation call and return extracted text plus raw response."""
    resp = _genai.models.generate_content(
        model=LLM_MODEL,
        contents=[prompt],
        config=generation_config,
    )
    return _extract_generated_text(resp), resp


def gen_text(prompt: str) -> str:
    """Generate text using LLM."""
    try:
        json_generation_config = genai_types.GenerateContentConfig(
            temperature=0.0,
            max_output_tokens=2192,
            response_mime_type="application/json",
        )
        generated_text, resp = _generate_with_config(
            prompt,
            json_generation_config,
        )
        if not generated_text:
            logging.warning(
                "LLM returned empty JSON-mode content for prompt length %s using model %s. %s",
                len(prompt),
                LLM_MODEL,
                _summarize_generation_response(resp),
            )
            text_generation_config = genai_types.GenerateContentConfig(
                temperature=0.0,
                max_output_tokens=2192,
            )
            generated_text, fallback_resp = _generate_with_config(
                prompt,
                text_generation_config,
            )
            if generated_text:
                logging.info(
                    "Recovered model output for prompt length %s using plain-text fallback mode.",
                    len(prompt),
                )
                return generated_text

            logging.warning(
                "LLM returned empty plain-text fallback content for prompt length %s using model %s. %s",
                len(prompt),
                LLM_MODEL,
                _summarize_generation_response(fallback_resp),
            )
        return generated_text
    except Exception as e:
        logging.error(f"LLM generation failed: {e}")
        return ""


def parse_json(s: str) -> Dict[str, Any]:
    """Parse JSON from string, handling markdown code blocks."""
    s = (s or "").strip()
    if s.startswith("```json"):
        s = s[7:]
    if s.startswith("```"):
        s = s[3:]
    if s.endswith("```"):
        s = s[:-3]
    s = s.strip()
    try:
        return json.loads(s)
    except Exception as e:
        logging.warning(f"parse_json failed: {e} | raw {s[:200]}")
        return {}


def parse_json_with_error(s: str) -> Tuple[Dict[str, Any], Optional[str], str]:
    """Parse JSON and preserve parser failures as structured errors."""
    cleaned = (s or "").strip()
    if cleaned.startswith("```json"):
        cleaned = cleaned[7:]
    if cleaned.startswith("```"):
        cleaned = cleaned[3:]
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3]
    cleaned = cleaned.strip()
    if not cleaned:
        return {}, "empty model response", cleaned
    try:
        parsed = json.loads(cleaned)
    except Exception as exc:
        extracted = _extract_json_object_text(cleaned)
        if extracted and extracted != cleaned:
            try:
                parsed = json.loads(extracted)
            except Exception:
                parsed = _parse_python_style_object(extracted)
            else:
                cleaned = extracted
        else:
            parsed = _parse_python_style_object(cleaned)
        if not isinstance(parsed, dict):
            return {}, f"JSON parse failed: {exc}", cleaned
    if not isinstance(parsed, dict):
        return {}, f"Expected JSON object, got {type(parsed).__name__}", cleaned
    return parsed, None, cleaned

def _extract_json_object_text(text_in: str) -> str:
    """Extract the largest JSON object substring from the input text."""
    start = text_in.find("{")
    end = text_in.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return ""
    return text_in[start : end + 1].strip()

def _parse_python_style_object(text_in: str) -> Dict[str, Any]:
    """Attempt to parse a Python-style dict object as a fallback."""
    try:
        parsed = ast.literal_eval(text_in)
    except Exception:
        return {}
    return parsed if isinstance(parsed, dict) else {}

def generate_structured_json(prompt: str, response_kind: str) -> Tuple[Dict[str, Any], Optional[str]]:
    """Generate structured JSON with a repair retry for non-empty malformed output."""
    raw_response = gen_text(prompt)
    parsed, error, cleaned = parse_json_with_error(raw_response)
    if error is None:
        return parsed, None

    if error == "empty model response":
        logging.warning("%s returned empty model response on first attempt.", response_kind)
        return {
            "status": "ERROR",
            "quality": {},
            "justification": (
                f"Validator received an empty model response for {response_kind}. "
                "Review prompt size, model availability, and logs."
            ),
            "citations": {},
            "validator_error": error,
            "raw_response_excerpt": "",
        }, error
    else:
        logging.warning("%s JSON parse failed on first attempt: %s", response_kind, error)

    repair_prompt = (
        f"{prompt}\n\n"
        "Your previous response was not valid JSON. "
        "Return STRICT JSON ONLY that matches the requested schema and do not include markdown fences.\n\n"
        "Previous invalid response:\n"
        f"{cleaned[:4000]}"
    )
    repaired_response = gen_text(repair_prompt)
    repaired, repair_error, repaired_cleaned = parse_json_with_error(repaired_response)
    if repair_error is None:
        return repaired, None

    combined_error = f"{error}; retry failed: {repair_error}"
    logging.error(f"{response_kind} JSON parse failed after retry: {combined_error}")
    return {
        "status": "ERROR",
        "quality": {},
        "justification": (
            f"Validator could not parse model output for {response_kind}. "
            "Review logs and retrieved evidence context."
        ),
        "citations": {},
        "validator_error": combined_error,
        "raw_response_excerpt": repaired_cleaned[:500],
    }, combined_error

def generate_structured_json_with_compact_retry(
    prompt: str,
    compact_retry_prompt: str,
    response_kind: str,
) -> Tuple[Dict[str, Any], Optional[str]]:
    """Generate structured JSON with a second, more compact retry prompt."""
    parsed, error = generate_structured_json(prompt, response_kind=response_kind)
    if error is None:
        return parsed, None

    if error == "empty model response":
        logging.warning(
            "%s standard prompt returned empty output; trying compact prompt length=%s.",
            response_kind,
            len(compact_retry_prompt),
        )
    else:
        logging.warning(
            "%s failed after standard retry; trying compact prompt length=%s.",
            response_kind,
            len(compact_retry_prompt),
        )
    compact_parsed, compact_error = generate_structured_json(
        compact_retry_prompt,
        response_kind=f"{response_kind} (compact retry)",
    )
    return compact_parsed, compact_error


def generate_architecture_summary_json(
    template_pdf_excerpts: str,
    evidence_snippets: List[str],
    diagram_refs: Optional[List[Dict[str, str]]] = None,
) -> Tuple[Dict[str, Any], Optional[str]]:
    """Generate architecture summary JSON with compact retry handling."""
    diagram_refs = diagram_refs or []
    diagram_block = "\n".join(
        _truncate_text(
            f"{ref.get('caption', '')}: {ref.get('doc_uri', '')}".strip(" :"),
            240,
        )
        for ref in diagram_refs[:4]
        if ref.get("caption") or ref.get("doc_uri")
    ) or "(None)"

    primary_snippets = compact_snippets(
        evidence_snippets,
        max_snippets=8,
        max_chars_per_snippet=900,
        max_total_chars=6000,
    )
    primary_prompt = ARCHITECTURE_SUMMARY_PROMPT.format(
        template_pdf_excerpts=_truncate_text(template_pdf_excerpts, 1800)
        or "(No template PDF excerpts)",
        evidence="\n\n---\n\n".join(primary_snippets) if primary_snippets else "(None)",
        diagrams=diagram_block,
    )

    retry_snippets = compact_snippets(
        evidence_snippets,
        max_snippets=4,
        max_chars_per_snippet=450,
        max_total_chars=1800,
    )
    retry_prompt = ARCHITECTURE_SUMMARY_PROMPT.format(
        template_pdf_excerpts=_truncate_text(template_pdf_excerpts, 1200)
        or "(No template PDF excerpts)",
        evidence="\n\n---\n\n".join(retry_snippets) if retry_snippets else "(None)",
        diagrams="\n".join(diagram_block.splitlines()[:2]) or "(None)",
    )

    logging.info(
        "Generating architecture summary JSON with primary prompt length=%s and compact prompt length=%s.",
        len(primary_prompt),
        len(retry_prompt),
    )
    return generate_structured_json_with_compact_retry(
        primary_prompt,
        retry_prompt,
        response_kind="Architecture Summary",
    )
# ============================================================================
# Document Processing
# ============================================================================

def read_pdf_text(pdf_path: str) -> str:
    """Extract text from PDF file."""
    pypdf_pages: List[str] = []
    fitz_pages: List[str] = []

    try:
        reader = PdfReader(pdf_path)
        pypdf_pages = [ (pg.extract_text() or "") for pg in reader.pages]
    except Exception as exc:
        logging.warning(f"PyPDF2 failed to extract text from '{pdf_path}': {exc}")

    try:
        import fitz  # PyMuPDF
        fitz_doc = fitz.open(pdf_path)
        try:
            fitz_pages = [(page.get_text("text") or "") for page in fitz_doc]
        finally:
            fitz_doc.close()
    except Exception as exc:
        if not pypdf_pages:
            logging.warning(f"PyMuPDF failed to extract text from '{pdf_path}': {exc}")

    if not pypdf_pages and not fitz_pages:
        logging.error(f"Both PyPDF2 and PyMuPDF failed to extract text from '{pdf_path}'. Returning empty string.")
        return ""

    merged_pages: List[str] = []
    for index in range(max(len(pypdf_pages), len(fitz_pages))):
        pypdf_text = pypdf_pages[index] if index < len(pypdf_pages) else ""
        fitz_text = fitz_pages[index] if index < len(fitz_pages) else ""
        chosen = fitz_text if len(fitz_text.strip()) > len(pypdf_text.strip()) else pypdf_text
        merged_pages.append(chosen)
    return "\n".join(merged_pages)


def read_docx_text(docx_path: str) -> str:
    """Extract text from DOCX file."""
    doc = DocxDocument(docx_path)
    parts: List[str] = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def compute_doc_hash(file_bytes: bytes) -> str:
    """Compute SHA-256 hash of file bytes."""
    return hashlib.sha256(file_bytes).hexdigest()


def split_text_local(
    text_in: str,
    chunk_size: int = 900,
    chunk_overlap: int = 90,
) -> List[str]:
    """Split text into chunks locally."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    return splitter.split_text(text_in)

def _truncate_text(value: str, limit: int) -> str:
    """Truncate text to a maximum character limit."""
    text = (value or "").strip()
    if limit <= 0 or len(text) <= limit:
        return text
    return text[: max(0, limit - 3)].rstrip() + "..."


def compact_snippets(
    snippets: List[str],
    max_snippets: int,
    max_chars_per_snippet: int,
    max_total_chars: int,
) -> List[str]:
    """Keep a unique, bounded set of snippets for prompt safety."""
    compacted: List[str] = []
    total_chars = 0
    seen = set()

    for snippet in snippets:
        normalized = _normalize_match_text(snippet)
        if not normalized or normalized in seen:
            continue

        seen.add(normalized)
        trimmed = _truncate_text(snippet, max_chars_per_snippet)
        if not trimmed:
            continue

        projected_total = total_chars + len(trimmed)
        if compacted and projected_total > max_total_chars:
            break

        compacted.append(trimmed)
        total_chars = projected_total

        if len(compacted) >= max_snippets or total_chars >= max_total_chars:
            break

    return compacted


_TEMPLATE_PDF_INDEX_CACHE: Dict[str, Tuple[List[str], np.ndarray]] = {}


def _get_template_pdf_index(pdf_text: str) -> Tuple[List[str], np.ndarray]:
    """Build and cache a reusable embedding index for the template PDF."""
    if not pdf_text.strip():
        return [], np.empty((0, 0), dtype=np.float32)

    cache_key = hashlib.sha256(pdf_text.encode("utf-8")).hexdigest()
    cached = _TEMPLATE_PDF_INDEX_CACHE.get(cache_key)
    if cached is not None:
        return cached

    chunks = split_text_local(pdf_text)
    if not chunks:
        cached = ([], np.empty((0, 0), dtype=np.float32))
    else:
        logging.info("Building template PDF embedding index for %s chunk(s).", len(chunks))
        cached = (
            chunks,
            np.array(
                embed_texts(chunks, task_type="RETRIEVAL_DOCUMENT", out_dim=EMBED_DIM),
                dtype=np.float32,
            ),
        )

    _TEMPLATE_PDF_INDEX_CACHE[cache_key] = cached
    return cached


def _get_bool_flag(node: Dict[str, Any], *keys: str) -> bool:
    """Read a boolean-like template flag from multiple possible key spellings."""
    for key in keys:
        if key in node:
            return bool(node.get(key))
    return False


def build_compact_prompt_block(
    snippets: List[str],
    empty_value: str,
    max_snippets: int,
    max_chars_per_snippet: int,
    max_total_chars: int,
    separator: str = "\n\n---\n\n",
) -> str:
    """Build a bounded prompt block from snippets."""
    compacted = compact_snippets(
        snippets,
        max_snippets=max_snippets,
        max_chars_per_snippet=max_chars_per_snippet,
        max_total_chars=max_total_chars,
    )
    return separator.join(compacted) if compacted else empty_value


def semantic_tag_text(tag: str) -> str:
    """Extract semantic content from tag (remove numbering)."""
    text = (tag or "").strip()
    # Remove leading numbers like "1.2.3"
    text = re.sub(r"^\s*\d+(?:[.,]\d+)*(?:\.[a-z])?\s+", "", text, flags=re.IGNORECASE)
    # Remove leading section labels like "A1", "B2.3"
    text = re.sub(r"^\s*[A-Za-z]\d+(?:[.,]\d+)*(?:\.[a-z])?\s+", "", text, flags=re.IGNORECASE)
    return text.strip() or (tag or "").strip()


def extract_tag_prefix(tag: str) -> str:
    """Extract numbering prefix from tag (e.g., '1.2' from '1.2 Title')."""
    match = re.match(
        r"^\s*([A-Za-z]?\d+(?:[.,]\d+)*(?:\.[a-z])?)\b",
        (tag or "").strip(),
        flags=re.IGNORECASE,
    )
    return match.group(1).strip().lower() if match else ""


def _normalize_match_text(text_in: str) -> str:
    """Normalize text for matching: lowercase, remove non-alphanumeric."""
    return re.sub(
        r"\s+",
        "",
        re.sub(r"[^a-z0-9]+", "", (text_in or "").lower()),
    ).strip()


def _normalize_line_text(text_in: str) -> str:
    """Normalize text while preserving word boundaries for line-based matching."""
    return re.sub(
        r"\s+",
        " ",
        re.sub(r"[^a-z0-9\s]+", " ", (text_in or "").lower()),
    ).strip()


def _normalize_multiline_text(text_in: str) -> str:
    """Normalize text line by line while preserving line boundaries."""
    lines = []
    for raw_line in re.split(r"\r?\n", text_in or ""):
        normalized_line = _normalize_line_text(raw_line)
        if normalized_line:
            lines.append(normalized_line)
    return "\n".join(lines)


def _tokenize_match_text(text_in: str) -> List[str]:
    """Tokenize text for loose semantic overlap checks."""
    return re.findall(r"[a-z0-9]+", (text_in or "").lower())


def _load_template_nodes(template_json_path: str) -> List[Dict[str, Any]]:
    """Load template nodes from either a top-level list or an object with children."""
    with open(template_json_path, "r", encoding="utf-8") as f:
        template = json.load(f)

    if isinstance(template, list):
        return [node for node in template if isinstance(node, dict)]

    if isinstance(template, dict):
        children = template.get("children")
        if isinstance(children, list):
            return [node for node in children if isinstance(node, dict)]

    raise ValueError(
        "Template JSON must be either a list of intent nodes or an object with a 'children' list."
    )


def _iter_requirement_nodes(
    nodes: List[Dict[str, Any]],
    current_intent: str = "",
) -> List[Tuple[str, Dict[str, Any]]]:
    """Flatten template trees into individual requirement nodes with inherited intent."""
    flattened: List[Tuple[str, Dict[str, Any]]] = []
    for node in nodes:
        if not isinstance(node, dict):
            continue
        intent_name = (node.get("intent") or current_intent or "").strip()
        tag = (node.get("tag") or "").strip()
        has_requirement_fields = any(key in node for key in ("musthave","mustHave", "goodToHave","goodTohave", "notes", "conditions", "keywords"))
        if tag and has_requirement_fields:
            flattened.append((intent_name, node))
        children = node.get("children")
        if isinstance(children, list) and children:
            flattened.extend(_iter_requirement_nodes(children, intent_name))
    return flattened


def _resolve_validation_scope_label(
    nar_id: str,
    release_number: str,
    rtype: str,
    explicit_doc_id: Optional[str],
) -> str:
    """Return a user-friendly validation scope summary."""
    if explicit_doc_id:
        return explicit_doc_id

    engine = get_engine()
    sql = text(f"""
        SELECT COUNT(DISTINCT doc_id) AS doc_count,
               MIN(doc_id) AS min_doc_id,
               MAX(doc_id) AS max_doc_id
        FROM {TABLE_EVIDENCE}
        WHERE nar_id=:nar AND release_number=:rel AND rtype=:rtype
    """)
    with engine.connect() as conn:
        row = conn.execute(sql, {"nar": nar_id, "rel": release_number, "rtype": rtype}).first()

    doc_count = int(row[0] or 0) if row else 0
    min_doc_id = row[1] if row else None
    max_doc_id = row[2] if row else None
    if doc_count <= 0:
        return "NO_MATCHING_DOCUMENTS"
    if doc_count == 1 and min_doc_id:
        return str(min_doc_id)
    if min_doc_id and max_doc_id and min_doc_id == max_doc_id:
        return str(min_doc_id)
    return f"ALL_MATCHING_DOCUMENTS ({doc_count} doc_ids)"


def _build_architecture_fallback(
    arch_snips: List[str],
    architecture_error: Optional[str],
    diagram_refs: Optional[List[Dict[str, str]]] = None,
) -> Dict[str, Any]:
    """Build a deterministic architecture summary when the model returns unusable output."""
    diagram_refs = diagram_refs or []
    diagram_text = " ".join(f"{ref.get('caption','')} {ref.get('doc_uri','')}" for ref in diagram_refs)
    evidence_text = _normalize_match_text(" ".join(arch_snips) + " " + diagram_text)

    def _has_any(*tokens: str) -> bool:
        return any(token in evidence_text for token in tokens)

    conceptual_present = _has_any("conceptual", "contextdiagram", "highlevelarchitecture")
    logical_present = _has_any("logical", "component", "interface", "service", "workflow")
    physical_present = _has_any("physical", "deployment", "gke", "cloudsql", "subnet", "firewall", "network")

    recommendations: List[str] = []
    if not conceptual_present:
        recommendations.append("Add a conceptual architecture view that explains the major business/system domains.")
    if not logical_present:
        recommendations.append("Add a logical architecture view with components, interfaces, and key data flows.")
    if not physical_present:
        recommendations.append("Add a physical or deployment architecture view with infrastructure placement details.")
    if architecture_error:
        recommendations.append("Review validator logs because the model did not return structured architecture JSON.")

    if arch_snips or diagram_refs:
        summary = (
            f"Architecture summary fallback used because the model response was unavailable. "
            f"Retrieved {len(arch_snips)} architecture evidence snippet(s) for manual review."
            "The model likely recieved more architecture context than it could reliably returna s strict JSON"
        )
    else:
        summary = (
            "Architecture summary fallback used because the model response was unavailable and no architecture evidence snippets were retrieved."
        )

    return {
        "summary": summary,
        "conceptual_present": conceptual_present,
        "logical_present": logical_present,
        "physical_present": physical_present,
        "label_alignment_issues": [],
        "diagram_refs": [
            _truncate_text(
                f"{ref.get('caption', '')}: {ref.get('doc_uri', '')}".strip(" :"),
                180,
            )
            for ref in diagram_refs[:3]
            if ref.get("caption") or ref.get("doc_uri")
        ],
        "recommendations": recommendations,
        "validator_error": architecture_error,
    }

def _build_topic_validation_fallback(
    tag: str,
    semantic_tag: str,
    keywords: List[str],
    evidence_snippets: List[str],
    guidance_snippets: List[str],
    template_pdf_snippets: List[str],
    validation_error: Optional[str],
    is_diagram: bool = False,
    diagram_refs: Optional[List[Dict[str, str]]] = None,
) -> Dict[str, Any]:
    """Build a deterministic topic-level fallback when the model returns unusable output."""
    diagram_refs = diagram_refs or []
    topic_match = semantic_topic_match_strength(
        semantic_tag, keywords, evidence_snippets
    )
    evidence_present = bool(evidence_snippets)
    diagram_present = bool(diagram_refs)

    if is_diagram:
        if diagram_present or topic_match in ("exact", "partial"):
            status = "PRESENT"
        elif evidence_present:
            status = "PARTIAL"
        else:
            status = "MISSING"

        quality = {
            "completeness": 4 if diagram_present else (3 if evidence_present else 0),
            "specificity": 4 if diagram_present else (2 if evidence_present else 0),
            "diagram_alignment": 4 if diagram_present else (2 if evidence_present else 0),
        }

        justification = (
            f"Used deterministic fallback for '{tag}' because the model returned "
            f"unusable output. Diagram references found: {len(diagram_refs)}; "
            f"evidence snippets found: {len(evidence_snippets)}."
        )

        citations = {
            "diagram": [
                _truncate_text(ref.get("caption", ""), 180)
                for ref in diagram_refs[:2]
            ],
            "text": [
                _truncate_text(snippet, 180)
                for snippet in evidence_snippets[:2]
            ],
            "template_pdf": [
                _truncate_text(snippet, 180)
                for snippet in template_pdf_snippets[:1]
            ],
        }
    else:
        if topic_match in ("exact", "partial"):
            status = "PRESENT"
        elif evidence_present:
            status = "PARTIAL"
        else:
            status = "MISSING"

        quality = {
            "completeness": 4 if topic_match == "exact" else (3 if evidence_present else 0),
            "specificity": 4 if topic_match == "exact" else (2 if evidence_present else 0),
            "traceability": 3 if evidence_present else 0,
        }

        justification = (
            f"Used deterministic fallback for '{tag}' because the model returned "
            f"unusable output. Evidence snippets found: {len(evidence_snippets)}; "
            f"semantic match={topic_match or 'none'}."
        )

        citations = {
            "evidence": [
                _truncate_text(snippet, 180)
                for snippet in evidence_snippets[:2]
            ],
            "guidance": [
                _truncate_text(snippet, 180)
                for snippet in guidance_snippets[:1]
            ],
            "template_pdf": [
                _truncate_text(snippet, 180)
                for snippet in template_pdf_snippets[:1]
            ],
        }

    return {
        "status": status,
        "quality": quality,
        "justification": justification,
        "citations": citations,
        "validator_error": validation_error,
    }

def topic_alignment_status(
    tag: str,
    semantic_tag: str,
    evidence_snippets: List[str],
) -> Optional[str]:
    """Check if topic is aligned with expected template tag."""
    tag_prefix = extract_tag_prefix(tag)
    if not tag_prefix or not evidence_snippets:
        return None
    normalized_topic = _normalize_line_text(semantic_tag)
    if not normalized_topic:
        return None
    tag_prefix_pattern = re.escape(tag_prefix).replace(r"\.", r"[.\s]+")
    topic_pattern = re.escape(normalized_topic).replace(r"\ ", r"\s+")
    aligned_pattern = re.compile(
        rf"(^|[\r\n])\s*({tag_prefix_pattern})(?:[\s:\-]+)({topic_pattern})\b",
        flags=re.IGNORECASE,
    )
    topic_line_pattern = re.compile(
        rf"(^|[\r\n])\s*(?:[A-Za-z]?\d+(?:[.\s]+\d+)*(?:\.[a-z])?)(?:[\s:\-]+){topic_pattern}\b",
        flags=re.IGNORECASE,
    )
    normalized_joined = "\n".join(
        _normalize_multiline_text(snippet) for snippet in evidence_snippets if snippet
    )
    if not normalized_joined.strip():
        return None
    if aligned_pattern.search(normalized_joined):
        return "aligned"
    if topic_line_pattern.search(normalized_joined):
        return "misaligned"
    return None


def build_topic_queries(
    intent_name: str,
    semantic_tag: str,
    tag: str,
    notes: str,
    conds: str,
    kw_text: str,
    requirement_text: str = "",
) -> List[str]:
    """Build multiple query strings for topic retrieval."""
    candidates = [
        " ".join(s for s in [semantic_tag, kw_text] if s).strip(),
        " ".join(s for s in [semantic_tag, requirement_text] if s).strip(),
        " ".join(s for s in [intent_name, semantic_tag, kw_text] if s).strip(),
        " ".join(s for s in [intent_name, semantic_tag, requirement_text] if s).strip(),
        " ".join(s for s in [semantic_tag, notes] if s).strip(),
        " ".join(s for s in [semantic_tag, conds] if s).strip(),
        " ".join(s for s in [tag, semantic_tag, notes, conds, kw_text, requirement_text] if s).strip(),
    ]
    queries: List[str] = []
    seen = set()
    for candidate in candidates:
        normalized = _normalize_match_text(candidate)
        if normalized and normalized not in seen:
            seen.add(normalized)
            queries.append(candidate)
    return queries


def build_topic_search_phrases(
    semantic_tag: str,
    tag: str,
    keywords: List[str],
    requirement_phrases: Optional[List[str]] = None,
) -> List[str]:
    """Build search phrases for keyword-based retrieval."""
    requirement_phrases = requirement_phrases or []
    candidates = [tag, semantic_tag] + [k for k in keywords if isinstance(k, str)] + requirement_phrases
    phrases: List[str] = []
    seen = set()
    for candidate in candidates:
        raw = (candidate or "").strip()
        normalized = _normalize_match_text(candidate)
        if raw and normalized and len(normalized) > 4 and normalized not in seen:
            seen.add(normalized)
            phrases.append(raw)
    return phrases


def retrieve_keyword_evidence_snippets(
    phrases: List[str],
    nar_id: str,
    release_number: str,
    rtype: str,
    doc_id: Optional[str],
    top_n: int,
) -> List[str]:
    """Retrieve evidence snippets using keyword matching."""
    valid_phrases = []
    seen = set()
    for phrase in phrases:
        normalized = _normalize_match_text(phrase)
        if normalized and normalized not in seen:
            seen.add(normalized)
            valid_phrases.append((phrase, normalized))
    if not valid_phrases:
        return []
    where = ["nar_id=:nar", "release_number=:rel", "rtype=:rtype"]
    params: Dict[str, Any] = {
        "nar": nar_id,
        "rel": release_number,
        "rtype": rtype,
        "topn": top_n,
    }
    if doc_id:
        where.append("doc_id=:doc_id")
        params["doc_id"] = doc_id
    phrase_conditions = []
    exact_checks = []
    for idx, (raw_phrase, normalized_phrase) in enumerate(valid_phrases):
        like_key = f"phrase_{idx}"
        exact_key = f"exact_{idx}"
        params[like_key] = f"%{raw_phrase}%"
        params[exact_key] = f"%{normalized_phrase}%"
        phrase_conditions.append(f"chunk ILIKE :{like_key}")
        exact_checks.append(
            f"CASE WHEN lower(regexp_replace(chunk, '[^a-zA-Z0-9]', '', 'g')) LIKE :{exact_key} THEN 0 ELSE 1 END"
        )
    sql = text(f"""
        SELECT chunk
        FROM {TABLE_EVIDENCE}
        WHERE {" AND ".join(where)}
          AND ({" OR ".join(phrase_conditions)})
        ORDER BY {", ".join(exact_checks)}, length(chunk)
        LIMIT :topn
    """)
    engine = get_engine()
    with engine.connect() as conn:
        rows = conn.execute(sql, params).fetchall()
    return [r[0] for r in rows if r and (r[0] or "").strip()]


def retrieve_topic_evidence_snippets(
    queries: List[str],
    nar_id: str,
    release_number: str,
    rtype: str,
    doc_id: Optional[str],
    top_n: int,
    search_phrases: Optional[List[str]] = None,
) -> List[str]:
    """Retrieve evidence snippets using semantic + keyword search."""
    merged: List[str] = []
    per_query_top_n = max(3, min(top_n, 4))
    seen = set()
    if search_phrases:
        lexical_hits = retrieve_keyword_evidence_snippets(
            search_phrases, nar_id, release_number, rtype, doc_id, top_n
        )
        for snippet in lexical_hits:
            normalized = _normalize_match_text(snippet)
            if normalized and normalized not in seen:
                seen.add(normalized)
                merged.append(snippet)
        if len(merged) >= top_n:
            return merged
    for query in queries:
        snippets = retrieve_evidence_snippets(
            query, nar_id, release_number, rtype, doc_id, per_query_top_n
        )
        for snippet in snippets:
            normalized = _normalize_match_text(snippet)
            if normalized and normalized not in seen:
                seen.add(normalized)
                merged.append(snippet)
        if len(merged) >= top_n:
            return merged
    return merged


def semantic_topic_match_strength(
    semantic_tag: str,
    keywords: List[str],
    evidence_snippets: List[str],
) -> Optional[str]:
    """Assess semantic match strength: 'exact', 'partial', or None."""
    if not evidence_snippets:
        return None
    joined = _normalize_match_text(" ".join(evidence_snippets))
    if not joined:
        return None
    candidate_phrases = [semantic_tag] + [k for k in keywords if isinstance(k, str)]
    normalized_phrases = []
    for phrase in candidate_phrases:
        normalized = _normalize_match_text(phrase)
        if normalized and len(normalized) >= 4:
            normalized_phrases.append(normalized)
    if any(phrase in joined for phrase in normalized_phrases):
        return "exact"
    stop_words = {
        "and", "the", "for", "with", "from", "into", "that", "this",
        "have", "must", "good", "include", "table", "section", "application",
        "overview", "details", "link", "keep", "applicable",
    }
    evidence_token_set = set(_tokenize_match_text(" ".join(evidence_snippets)))
    semantic_tokens : List[str] = []
    for phrase in [semantic_tag] + [k for k in keywords if isinstance(k, str)]:
        for token in _tokenize_match_text(phrase):
            if len(token) > 4 and token not in stop_words and token not in semantic_tokens:
                semantic_tokens.append(token)
    if not semantic_tokens:
        return None
    token_hits = sum(
        1
        for token in semantic_tokens
        if token in evidence_token_set
    )
    required_hits = min(len(semantic_tokens), max(2, (len(semantic_tokens) + 1) // 2))
    return "partial" if token_hits >= required_hits else None

def extract_requirement_phrases(notes: str, conds: str) -> List[str]:
    """Extract concise requirement cues from notes/conditions for retrieval and matching."""
    phrases: List[str] = []
    seen = set()

    def add_phrase(value: str) -> None:
        raw = (value or "").strip(" .;:-")
        normalized = _normalize_match_text(raw)
        if not raw or len(normalized) < 5 or normalized in seen:
            return
        seen.add(normalized)
        phrases.append(raw)

    include_match = re.search(
        r"include\s*:\s*(.+?)(?:\.|$)",
        notes or "",
        flags=re.IGNORECASE,
    )
    if include_match:
        include_block = include_match.group(1)
        for part in re.split(r",|;|\band\b", include_block, flags=re.IGNORECASE):
            add_phrase(part)

    for text_block in [notes, conds]:
        for part in re.split(r"[.;]\s+", text_block or ""):
            if len(_normalize_match_text(part)) >= 10:
                add_phrase(part)

    return phrases[:6]


def should_promote_semantic_match(
    topic_match: Optional[str],
    topic_alignment: Optional[str],
    keywords: List[str],
    evidence_snippets: List[str],
) -> bool:
    """Allow semantic promotion only when the evidence signal is strong enough."""
    if topic_match == "exact":
        return True

    if topic_match != "partial":
        return False

    if topic_alignment in ("aligned", "misaligned"):
        return True

    evidence_count = sum(
        1 for snippet in evidence_snippets if (snippet or "").strip()
    )
    if evidence_count < 2:
        return False

    normalized_joined = _normalize_match_text("\n".join(evidence_snippets))
    strong_keyword_hits = 0

    for keyword in keywords:
        normalized_keyword = _normalize_match_text(keyword)
        if (
            normalized_keyword
            and len(normalized_keyword) >= 5
            and normalized_keyword in normalized_joined
        ):
            strong_keyword_hits += 1

    return strong_keyword_hits >= 2

def pick_best_template_snippets_for_tag(
    tag_query: str,
    pdf_text: str,
    top_n: int = 3,
) -> List[str]:
    """Pick best matching snippets from template PDF using embeddings."""
    chunks, emb_chunks = _get_template_pdf_index(pdf_text)
    if not chunks or emb_chunks.size == 0:
        return []
    qv = np.array(embed_query_text(tag_query, out_dim=EMBED_DIM), dtype=np.float32).reshape(1, -1)
    sims = cosine_similarity(qv, emb_chunks)[0]
    idx = np.argsort(-sims)[:top_n]
    return [chunks[i] for i in idx]


# ============================================================================
# Vector Literal & Retrieval Functions
# ============================================================================

def _vec_literal(vec: List[float]) -> str:
    """Format vector as PostgreSQL pgvector literal."""
    return "[" + ",".join(f"{x:.6f}" for x in vec) + "]"


def retrieve_evidence_snippets(
    query: str,
    nar_id: str,
    release_number: str,
    rtype: str,
    doc_id: Optional[str],
    top_n: int,
) -> List[str]:
    """Retrieve evidence snippets using semantic search."""
    q_emb = embed_query_text(query)
    vec = _vec_literal(q_emb)
    where = ["nar_id=:nar", "release_number=:rel", "rtype=:rtype"]
    params = {
        "nar": nar_id,
        "rel": release_number,
        "rtype": rtype,
        "vec": vec,
        "topn": top_n,
    }
    if doc_id:
        where.append("doc_id=:doc_id")
        params["doc_id"] = doc_id
    sql = text(f"""
        SELECT chunk
        FROM {TABLE_EVIDENCE}
        WHERE {" AND ".join(where)}
        ORDER BY embedding <-> CAST(:vec AS vector)
        LIMIT :topn
    """)
    engine = get_engine()
    with engine.connect() as conn:
        rows = conn.execute(sql, params).fetchall()
    return [r[0] for r in rows if r and (r[0] or "").strip()]


def retrieve_guidance_snippets(
    query: str,
    rtype: str,
    top_n: int = 3,
) -> List[str]:
    """Retrieve guidance snippets."""
    try:
        qv = embed_query_text(query)
        vec = _vec_literal(qv)
        sql = text(f"""
            SELECT chunk
            FROM {TABLE_GUIDANCE}
            WHERE rtype=:rtype
            ORDER BY embedding <-> CAST(:vec AS vector)
            LIMIT :topn
        """)
        engine = get_engine()
        with engine.connect() as c:
            rows = c.execute(sql, {"rtype": rtype, "vec": vec, "topn": top_n}).fetchall()
        return [r[0] for r in rows if r and r[0]]
    except Exception as e:
        logging.warning(f"guidance retrieval failed: {e}")
        return []


def retrieve_mm_diagrams(
    query_text: str,
    nar_id: str,
    release_number: str,
    rtype: str,
    doc_id: Optional[str],
    top_n: int = 3,
    dim: int = MM_DIM,
) -> List[Dict[str, str]]:
    """Retrieve multimodal diagram references."""
    qv = mm_embed_text(query_text, dim)
    vec = _vec_literal(qv)
    where = ["nar_id=:nar", "release_number=:rel", "rtype=:rtype"]
    params: Dict[str, Any] = {
        "nar": nar_id,
        "rel": release_number,
        "rtype": rtype,
        "vec": vec,
        "topn": top_n,
    }
    if doc_id:
        where.append("doc_id=:doc_id")
        params["doc_id"] = doc_id
    sql = text(f"""
        SELECT caption, doc_uri
        FROM {TABLE_MM_EVID}
        WHERE {" AND ".join(where)}
        ORDER BY embedding <-> CAST(:vec AS vector)
        LIMIT :topn
    """)
    engine = get_engine()
    with engine.connect() as c:
        rows = c.execute(sql, params).fetchall()
    return [{"caption": r[0], "doc_uri": r[1]} for r in rows]


def retrieve_mm_diagrams_safe(
    query_text: str,
    nar_id: str,
    release_number: str,
    rtype: str,
    doc_id: Optional[str],
    top_n: int = 3,
    dim: int = MM_DIM,
) -> List[Dict[str, str]]:
    """Best-effort multimodal retrieval that degrades to text-only validation."""
    try:
        return retrieve_mm_diagrams(
            query_text=query_text,
            nar_id=nar_id,
            release_number=release_number,
            rtype=rtype,
            doc_id=doc_id,
            top_n=top_n,
            dim=dim,
        )
    except Exception as exc:
        logging.warning(
            "Multimodal retrieval failed for nar_id=%s, release=%s, rtype=%s, query=%r: %s",
            nar_id,
            release_number,
            rtype,
            _truncate_text(query_text, 120),
            exc,
        )
        return []


# ============================================================================
# LLM PROMPTS
# ============================================================================

PRESENCE_QUALITY_PROMPT = """You are a validation agent.

Decide PRESENCE and score QUALITY using:
- Template tag & notes
- Template-PDF criteria excerpts (authoritative instructions)
- Optional Guidance snippets (org guidance)
- Evidence text snippets

Return STRICT JSON only.

Important evaluation rule:
Judge semantic content, not exact section numbering.
If the right content appears under a different numbering or heading label (for example 1.2 instead of 1.1), do NOT mark it PARTIAL or MISSING for that reason alone.
Only mark PARTIAL or MISSING when required content itself is incomplete, generic, contradictory, or absent.

TAG: {tag}
INTENT: {intent}
SEMANTIC TAG: {semantic_tag}
NOTES: {notes}
REQUIRED: {required_flag}

TEMPLATE PDF EXCERPTS:
{template_pdf_excerpts}

GUIDANCE SNIPPETS:
{guidance}

EVIDENCE SNIPPETS:
{evidence}

Respond JSON:
{{
  "status": "PRESENT" | "PARTIAL" | "MISSING" | "EMPTY",
  "quality": {{
    "completeness": 0-5,
    "specificity": 0-5,
    "traceability": 0-5
  }},
  "justification": "short one or two sentences",
  "citations": {{
    "evidence": ["short excerpt 1", "short excerpt 2"],
    "guidance": ["short excerpt or empty"],
    "template_pdf": ["short excerpt 1"]
  }}
}}
"""

DIAGRAM_PROMPT = """You are validating whether the required DIAGRAM exists and is of good quality.

Use Template tag/notes, Template-PDF criteria, optional Guidance, and DIAGRAM REFERENCES.
Return STRICT JSON only.

Important evaluation rule:
Judge semantic alignment, not exact section numbering.
If the diagram/content is present under a different numbered section, do NOT penalize that numbering mismatch by itself.

TAG: {tag}
INTENT: {intent}
SEMANTIC TAG: {semantic_tag}
NOTES: {notes}
REQUIRED: {required_flag}

TEMPLATE_PDF_EXCERPTS:
{template_pdf_excerpts}

GUIDANCE SNIPPETS:
{guidance}

EVIDENCE TEXT SNIPPETS:
{evidence}

DIAGRAM_REFERENCES (top-k by similarity):
{diagrams}

Respond JSON:
{{
  "status": "PRESENT" | "PARTIAL" | "MISSING" | "EMPTY",
  "quality": {{
    "completeness": 0-5,
    "specificity": 0-5,
    "diagram_alignment": 0-5
  }},
  "justification": "short one or two sentences",
  "citations": {{
    "diagram": ["<caption or uri>"],
    "text": ["short excerpt or empty"],
    "template_pdf": ["short excerpt 1"]
  }}
}}
"""

ARCHITECTURE_SUMMARY_PROMPT = """Summarize architecture coverage and quality (conceptual, logical, physical).
Use text evidence snippets, architecture diagram references and Template-PDF expectations.

TEMPLATE_PDF_EXCERPTS:
{template_pdf_excerpts}

EVIDENCE SNIPPETS:
{evidence}

ARCHITECTURE DIAGRAM REFRENCES:
{diagrams}

Return STRICT JSON:
{{
  "summary": "concise summary",
  "conceptual_present": true | false,
  "logical_present": true | false,
  "physical_present": true | false,
    "diagram_refs": ["short caption or uri"],
  "label_alignment_issues": ["example or empty"],
  "recommendations": ["short, actionable rec #1", "short, actionable rec #2"]
}}
"""


# ============================================================================
# Core Validation Engine
# ============================================================================

def run_validation(
    template_json_path: str,
    template_pdf_path: str,
    nar_id: str,
    application_name: str,
    release_number: str,
    rtype: str,
    top_k: int = 6,
    scope_doc_id: Optional[str] = None,
    scope_doc_hash: Optional[str] = None,
    enable_mm: bool = True,
) -> Dict[str, Any]:
    """Run full validation workflow."""
    effective_doc_id = (scope_doc_id or scope_doc_hash or "").strip() or None
    if scope_doc_id and scope_doc_hash and scope_doc_id != scope_doc_hash:
        raise ValueError("scope_doc_id and scope_doc_hash must match when both are provided.")

    if EMBED_DIM is None:
        logging.warning(
            f"Unknown embedding model '{EMBED_MODEL}'. "
            f"Ensure VECTOR(dim) matches model output."
        )

    template_nodes = _load_template_nodes(template_json_path)
    requirement_nodes = _iter_requirement_nodes(template_nodes)
    if not requirement_nodes:
        raise ValueError(f"No requirement nodes found in template JSON: {template_json_path}")
    logging.info(
        "Loaded %s requirement node(s): must-have=%s, good-to-have=%s.",
        len(requirement_nodes),
        sum(
            1
            for _, node in requirement_nodes
            if _get_bool_flag(node, "musthave", "mustHave")
        ),
        sum(
            1
            for _, node in requirement_nodes
            if _get_bool_flag(node, "goodToHave", "goodtohave", "good_to_have")
        ),
    )

    # Extract Template PDF text
    if template_pdf_path.lower().endswith(".pdf"):
        pdf_text = read_pdf_text(template_pdf_path)
    elif template_pdf_path.lower().endswith(".docx"):
        pdf_text = read_docx_text(template_pdf_path)
    else:
        logging.warning("Template PDF path is not PDF/DOCX; skipping PDF guidance extraction.")
        pdf_text = ""

    results: List[Dict[str, Any]] = []

    # Iterate over template
    for intent_name, child in requirement_nodes:
        tag = (child.get("tag") or "").strip()
        semantic_tag = semantic_tag_text(tag)
        mreq = _get_bool_flag(child, "musthave", "mustHave")
        greq = _get_bool_flag(child, "goodToHave", "goodtohave", "good_to_have")

        notes = (child.get("notes", "") or "").strip()
        conds = (child.get("conditions", "") or "").strip()

        kws = child.get("keywords") or []
        requirement_phrases = extract_requirement_phrases(notes, conds)
        combined_match_terms = [k for k in kws if isinstance(k, str)] + requirement_phrases

        kw_text = " ".join(
            k.strip()
            for k in combined_match_terms
            if isinstance(k, str) and k.strip()
        )

        topic_queries = build_topic_queries(
            intent_name,
            semantic_tag,
            tag,
            notes,
            conds,
            kw_text,
            requirement_text="\n".join(requirement_phrases),
        )
        topic_search_phrases = build_topic_search_phrases(
            semantic_tag,
            tag,
            kws,
            requirement_phrases=requirement_phrases,
        )
        query_full = topic_queries[0] if topic_queries else " ".join(
            s for s in [semantic_tag, kw_text] if s
        ).strip()

        ev_snips = retrieve_topic_evidence_snippets(
            topic_queries,
            nar_id,
            release_number,
            rtype,
            effective_doc_id,
            top_k,
            topic_search_phrases,
        )
        s_join = build_compact_prompt_block(
            ev_snips[:top_k],
            empty_value="(no snippets found)",
            max_snippets=4,
            max_chars_per_snippet=700,
            max_total_chars=2600,
        )

        g_snips = retrieve_guidance_snippets(query_full, rtype, top_n=3)
        g_join = build_compact_prompt_block(
            g_snips,
            empty_value="(no guidance context)",
            max_snippets=2,
            max_chars_per_snippet=500,
            max_total_chars=1000,
        )

        pdf_excerpts = pick_best_template_snippets_for_tag(
            query_full,
            pdf_text,
            top_n=3,
        )
        pdf_join = build_compact_prompt_block(
            pdf_excerpts,
            empty_value="(no template-pdf excerpt)",
            max_snippets=2,
            max_chars_per_snippet=500,
            max_total_chars=1000,
        )

        is_diagram = (
            ("diagram" in tag.lower())
            or ("diagram" in notes.lower())
            or ("diagram" in conds.lower())
            or any(isinstance(k, str) and "diagram" in k.lower() for k in kws)
        )

        mm_hits: List[Dict[str, str]] = []
        if enable_mm and is_diagram:
            mm_hits = retrieve_mm_diagrams_safe(
                query_full,
                nar_id,
                release_number,
                rtype,
                effective_doc_id,
                top_n=3,
                dim=MM_DIM,
            )
            diag_lines = [
                _truncate_text(f"{hit['caption']}: {hit['doc_uri']}", 260)
                for hit in mm_hits
            ]
            diag_block = "\n".join(line for line in diag_lines if line) or "(none)"

            diagram_prompt = DIAGRAM_PROMPT.format(
                tag=tag,
                intent=intent_name,
                semantic_tag=semantic_tag,
                notes=_truncate_text(notes, 500),
                required_flag=(
                    "MUST HAVE" if mreq else ("GOOD_TO_HAVE" if greq else "OPTIONAL")
                ),
                template_pdf_excerpts=pdf_join,
                guidance=g_join,
                evidence=s_join,
                diagrams=diag_block,
            )
            diagram_compact_retry_prompt = DIAGRAM_PROMPT.format(
                tag=tag,
                intent=intent_name,
                semantic_tag=semantic_tag,
                notes=_truncate_text(notes, 240),
                required_flag=(
                    "MUST HAVE" if mreq else ("GOOD_TO_HAVE" if greq else "OPTIONAL")
                ),
                template_pdf_excerpts=build_compact_prompt_block(
                    pdf_excerpts,
                    empty_value="(no template-pdf excerpt)",
                    max_snippets=1,
                    max_chars_per_snippet=280,
                    max_total_chars=280,
                ),
                guidance=build_compact_prompt_block(
                    g_snips,
                    empty_value="(no guidance context)",
                    max_snippets=1,
                    max_chars_per_snippet=220,
                    max_total_chars=220,
                ),
                evidence=build_compact_prompt_block(
                    ev_snips[:top_k],
                    empty_value="(no snippets found)",
                    max_snippets=2,
                    max_chars_per_snippet=380,
                    max_total_chars=760,
                ),
                diagrams="\n".join(diag_lines[:2]) or "(none)",
            )

            logging.info(
                "Diagram prompt prepared for tag '%s' with length=%s; compact retry length=%s.",
                tag,
                len(diagram_prompt),
                len(diagram_compact_retry_prompt),
            )
            p_json, validation_error = generate_structured_json_with_compact_retry(
                diagram_prompt,
                diagram_compact_retry_prompt,
                response_kind=f"diagram validation for tag '{tag}'",
            )

            if validation_error:
                logging.warning(
                    "Using deterministic diagram fallback for tag '%s' after model failure.",
                    tag,
                )
                p_json = _build_topic_validation_fallback(
                    tag=tag,
                    semantic_tag=semantic_tag,
                    keywords=combined_match_terms,
                    evidence_snippets=ev_snips,
                    guidance_snippets=g_snips,
                    template_pdf_snippets=pdf_excerpts,
                    validation_error=validation_error,
                    is_diagram=True,
                    diagram_refs=mm_hits,
                )
        else:
            presence_prompt = PRESENCE_QUALITY_PROMPT.format(
                tag=tag,
                intent=intent_name,
                semantic_tag=semantic_tag,
                notes=_truncate_text(notes, 500),
                required_flag=(
                    "MUST HAVE" if mreq else ("GOOD_TO_HAVE" if greq else "OPTIONAL")
                ),
                template_pdf_excerpts=pdf_join,
                guidance=g_join,
                evidence=s_join,
            )
            presence_compact_retry_prompt = PRESENCE_QUALITY_PROMPT.format(
                tag=tag,
                intent=intent_name,
                semantic_tag=semantic_tag,
                notes=_truncate_text(notes, 240),
                required_flag=(
                    "MUST HAVE" if mreq else ("GOOD_TO_HAVE" if greq else "OPTIONAL")
                ),
                template_pdf_excerpts=build_compact_prompt_block(
                    pdf_excerpts,
                    empty_value="(no template-pdf excerpt)",
                    max_snippets=1,
                    max_chars_per_snippet=280,
                    max_total_chars=280,
                ),
                guidance=build_compact_prompt_block(
                    g_snips,
                    empty_value="(no guidance context)",
                    max_snippets=1,
                    max_chars_per_snippet=220,
                    max_total_chars=220,
                ),
                evidence=build_compact_prompt_block(
                    ev_snips[:top_k],
                    empty_value="(no snippets found)",
                    max_snippets=2,
                    max_chars_per_snippet=380,
                    max_total_chars=760,
                ),
            )

            logging.info(
                "Presence prompt prepared for tag '%s' with length=%s; compact retry length=%s.",
                tag,
                len(presence_prompt),
                len(presence_compact_retry_prompt),
            )
            p_json, validation_error = generate_structured_json_with_compact_retry(
                presence_prompt,
                presence_compact_retry_prompt,
                response_kind=f"presence validation for tag '{tag}'",
            )

            if validation_error:
                logging.warning(
                    "Using deterministic presence fallback for tag '%s' after model failure.",
                    tag,
                )
                p_json = _build_topic_validation_fallback(
                    tag=tag,
                    semantic_tag=semantic_tag,
                    keywords=combined_match_terms,
                    evidence_snippets=ev_snips,
                    guidance_snippets=g_snips,
                    template_pdf_snippets=pdf_excerpts,
                    validation_error=validation_error,
                )

        status = (
            (p_json.get("status") if isinstance(p_json, dict) else None) or "MISSING"
        ).upper()
        justification = p_json.get("justification", "") if isinstance(p_json, dict) else ""
        result_quality = p_json.get("quality", {}) if isinstance(p_json, dict) else {}
        result_citations = p_json.get("citations", {}) if isinstance(p_json, dict) else {}
        result_validation_error = validation_error

        tag_1 = (tag or "").strip().lower()
        topic_match = semantic_topic_match_strength(
            semantic_tag,
            combined_match_terms,
            ev_snips,
        )
        topic_alignment = topic_alignment_status(tag, semantic_tag, ev_snips)
        promote_semantic_match = should_promote_semantic_match(
            topic_match=topic_match,
            topic_alignment=topic_alignment,
            keywords=combined_match_terms,
            evidence_snippets=ev_snips,
        )

        if promote_semantic_match and status in ("MISSING", "PARTIAL", "ERROR"):
            status = "PRESENT"
            if topic_alignment == "aligned":
                topic_note = (
                    f"Topic '{semantic_tag}' was found in the evidence and is well-aligned with the expected template tag '{tag}'. "
                    "Marked PRESENT based on strong semantic match and alignment, despite any numbering or labeling differences."
                )
            elif topic_alignment == "misaligned":
                topic_note = (
                    f"Topic '{semantic_tag}' appears in the evidence, but not under the expected template tag '{tag}'. "
                    "Marked PRESENT because the content is present but the section/tag alignment does not match the template."
                )
            else:
                topic_note = (
                    f"Topic '{semantic_tag}' was identified in the evidence with a partial semantic match. "
                    "Marked PRESENT based on the presence of relevant content, even though it may not fully meet all criteria or be perfectly aligned."
                )
            justification = topic_note
            if not isinstance(result_citations, dict):
                result_citations = {}

        if tag_1 in ("0.1 nar id", "8.1 nar id", "nar id") and nar_id:
            status = "PRESENT"
            justification = f"NAR ID '{nar_id}' was resolved from evidence."

        if tag_1 in (
            "0.2 app name",
            "0.2 application name",
            "app name",
            "application name",
        ) and application_name:
            status = "PRESENT"
            justification = f"Application name '{application_name}' was resolved from evidence."

        if tag_1 in ("0.3 r-type", "0.3 r type", "r-type", "r type", "rtype") and rtype:
            status = "PRESENT"
            justification = f"R-Type '{rtype}' was provided for validation."

        results.append({
            "intent": intent_name,
            "tag": tag,
            "musthave": mreq,
            "good_to_have": greq,
            "status": status,
            "quality": result_quality,
            "justification": justification,
            "citations": result_citations,
            "diagram_refs": mm_hits,
            "snippets": ev_snips[:top_k],
            "used_guidance": bool(g_snips),
            "template_pdf_used": bool(pdf_excerpts),
            "fallback_used": bool(validation_error),
            "match_strength": topic_match,
            "topic_alignment": topic_alignment,
            "evidence_snippet_count": len(ev_snips),
            "guidance_snippet_count": len(g_snips),
            "template_pdf_snippet_count": len(pdf_excerpts),
            "diagram_ref_count": len(mm_hits),
            "validator_error": result_validation_error,
        })

    # Architecture summary (conceptual, logical, physical)
    arch_query = "overall Architecture Conceptual Logical Physical GKE Cloud SQL Interfaces components"
    arch_snips = retrieve_evidence_snippets(
        arch_query, nar_id, release_number, rtype, effective_doc_id, top_n=18
    )
    arch_mm_hits = (
        retrieve_mm_diagrams_safe(
            "overall architecture conceptual logical physical deployment topology interfaces diagram",
            nar_id,
            release_number,
            rtype,
            effective_doc_id,
            top_n=4,
            dim=MM_DIM,
        )
        if enable_mm
        else []
    )

    arch_pdf = pick_best_template_snippets_for_tag(
        "Overall Architecture Diagrams Conceptual Logical Physical",
        pdf_text,
        top_n=3,
    )
    arch_pdf_join = "\n".join(arch_pdf) if arch_pdf else "(no template-pdf excerpt)"

    arch_json, architecture_error = generate_architecture_summary_json(
        template_pdf_excerpts=arch_pdf_join,
        evidence_snippets=arch_snips,
        diagram_refs=arch_mm_hits,
    )
    if architecture_error:
        arch_json = _build_architecture_fallback(arch_snips, architecture_error, arch_mm_hits)

    logging.info(f"arch_json result: {arch_json}")

    if not results:
        raise RuntimeError("No validation rows were produced from the template. Check template structure and traversal logic.")

    must_total = sum(1 for item in results if item.get("musthave"))
    must_met = sum(1 for item in results if item.get("musthave") and item.get("status") == "PRESENT")
    good_total = sum(1 for item in results if item.get("good_to_have"))
    good_met = sum(1 for item in results if item.get("good_to_have") and item.get("status") == "PRESENT")

    scope_label = _resolve_validation_scope_label(
        nar_id=nar_id,
        release_number=release_number,
        rtype=rtype,
        explicit_doc_id=effective_doc_id,
    )

    summary = {
        "nar_id": nar_id,
        "application_name": application_name,
        "release_number": release_number,
        "rtype": rtype,
        "doc_id_scope": scope_label,
        "doc_hash_scope": scope_label,
        "must_have_total": must_total,
        "must_have_coverage_pct": round((must_met / must_total * 100.0), 1) if must_total else 0.0,
        "must_have_met": must_met,
        "good_to_have_total": good_total,
        "good_to_have_coverage_pct": round((good_met / good_total * 100.0), 1) if good_total else 0.0,
        "fallback_topic_count": sum(1 for item in results if item.get("fallback_used")),
        "topics_with_validator_warnings": sum(1 for item in results if item.get("validator_error")),
        "architecture_evidence_snippet_count": len(arch_snips),
        "architecture_diagram_ref_count": len(arch_mm_hits),
        "architecture_template_excerpt_count": len(arch_pdf),
        "good_to_have_met": good_met,
        "architecture": arch_json,
        "architecture_error": architecture_error,
    }

    return {"summary": summary, "per_intent": results}


# ============================================================================
# PDF Generation
# ============================================================================

def generate_summary_pdf(
    result: Dict[str, Any],
    pdf_output_path: str = "summary_report.pdf",
) -> bool:
    """Generate summary PDF report."""
    try:
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4

        if not (pdf_output_path or "").strip():
            pdf_output_path = "summary_report.pdf"

        if not pdf_output_path.lower().endswith(".pdf"):
            pdf_output_path = f"{pdf_output_path}.pdf"

        styles = getSampleStyleSheet()
        link_style = styles["Normal"].clone("link_style")
        link_style.textColor = colors.blue
        link_style.underline = True

        def safe_pdf_text(value: Any) -> str:
            text = xml_escape(str(value or ""), {'"': "&quot;", "'": "&#39;"})
            return text.replace("\r\n", "\n").replace("\r", "\n").replace("\n", "<br/>")

        summary = result.get("summary", {})
        per_intent = result.get("per_intent", [])
        architecture = summary.get("architecture", {}) or {}

        nar_id = summary.get("nar_id", "")
        application_name = summary.get("application_name", "")
        release_number = summary.get("release_number", "")
        rtype = summary.get("rtype", "")
        must_have_total = summary.get("must_have_total", 0)
        must_have_met = summary.get("must_have_met", 0)
        must_have_coverage = summary.get("must_have_coverage_pct", 0.0)
        must_have_diff = must_have_total - must_have_met

        good_to_have_total = summary.get("good_to_have_total", 0)
        good_to_have_met = summary.get("good_to_have_met", 0)
        good_to_have_diff = good_to_have_total - good_to_have_met
        good_to_have_coverage = summary.get("good_to_have_coverage_pct", 0.0)

        architecture_summary = architecture.get("summary", "")
        if not architecture_summary:
            architecture_summary = (
                f"This report is for NAR ID {nar_id}, application {application_name}, "
                f"release {release_number}, and rtype {rtype}."
            )

        result_status = "PASS" if must_have_total == must_have_met else "FAIL"

        doc = SimpleDocTemplate(
            pdf_output_path,
            pagesize=A4,
            leftMargin=40,
            rightMargin=40,
            topMargin=40,
            bottomMargin=40,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=22,
            textColor=colors.HexColor("#1a5e80"),
            spaceAfter=18,
            alignment=1,
            fontName="Helvetica-Bold",
        )
        normal_style = styles["BodyText"]
        small_style = ParagraphStyle(
            "SmallStyle",
            parent=styles["BodyText"],
            fontSize=10,
            textColor=colors.HexColor("#333333"),
        )

        def section_heading(text):
            tbl = Table(
                [[Paragraph(f"<b>{text}</b>", ParagraphStyle(
                    "SectionHeadingText",
                    parent=styles["BodyText"],
                    fontSize=11,
                    textColor=colors.white,
                    fontName="Helvetica-Bold",
                    alignment=0,
                ))]],
                colWidths=[500],
            )
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#2d7fa3")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            return tbl

        story = []
        story.append(Paragraph("SDLC Control Validation Summary", title_style))
        story.append(Spacer(1, 8))
        story.append(section_heading("Document Information"))
        story.append(Spacer(1, 8))

        metadata_table_data = [
            [
                Paragraph("<b>NAR ID</b>", normal_style),
                Paragraph(safe_pdf_text(nar_id), normal_style),
            ],
            [
                Paragraph("<b>Application</b>", normal_style),
                Paragraph(safe_pdf_text(application_name), normal_style),
            ],
            [
                Paragraph("<b>Release</b>", normal_style),
                Paragraph(safe_pdf_text(release_number), normal_style),
            ],
            [
                Paragraph("<b>Type</b>", normal_style),
                Paragraph(safe_pdf_text(rtype), normal_style),
            ],
        ]
        metadata_table = Table(metadata_table_data, colWidths=[120, 380])
        metadata_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e8f1f5")),
            ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#f9f9f9")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ]))
        story.append(metadata_table)
        story.append(Spacer(1, 12))

        story.append(section_heading("Result"))
        story.append(Spacer(1, 6))

        reason_text = (
            "All must-have requirements are met."
            if result_status == "PASS"
            else f"{must_have_diff} must-have requirement(s) are not met."
        )
        result_line_style = ParagraphStyle(
            "ResultLineStyle",
            parent=styles["BodyText"],
            fontSize=11,
            textColor=colors.HexColor("#222222"),
            spaceAfter=6,
            leading=14,
        )
        status_color = "#2e7d32" if result_status == "PASS" else "#c62828"
        story.append(
            Paragraph(
                f'Result: <font color="{status_color}"><b>{result_status}</b></font>',
                result_line_style,
            )
        )
        story.append(Paragraph(f"Reason: {reason_text}", result_line_style))
        story.append(Spacer(1, 10))

        story.append(section_heading("Must-Have Requirements"))
        story.append(Spacer(1, 8))

        metrics_data = [
            [
                Paragraph("<b>Total</b>", small_style),
                Paragraph("<b>Met</b>", small_style),
                Paragraph("<b>Must-Have Not Met</b>", small_style),
                Paragraph("<b>Coverage</b>", small_style),
            ],
            [
                Paragraph(str(must_have_total), small_style),
                Paragraph(str(must_have_met), small_style),
                Paragraph(str(must_have_diff), small_style),
                Paragraph(f"<b>{must_have_coverage:.1f}%</b>", small_style),
            ],
        ]
        metrics_table = Table(metrics_data, colWidths=[80, 80, 160, 80])
        metrics_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#f9f9f9")),
            ("GRID", (0, 0), (-1, -1), 0.75, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(metrics_table)
        story.append(Spacer(1, 12))

        story.append(section_heading("Good-to-Have Requirements"))
        story.append(Spacer(1, 8))

        good_metrics_data = [
            [
                Paragraph("<b>Total</b>", small_style),
                Paragraph("<b>Met</b>", small_style),
                Paragraph("<b>Good-to-Have Not Met</b>", small_style),
                Paragraph("<b>Coverage</b>", small_style),
            ],
            [
                Paragraph(str(good_to_have_total), small_style),
                Paragraph(str(good_to_have_met), small_style),
                Paragraph(str(good_to_have_diff), small_style),
                Paragraph(f"<b>{good_to_have_coverage:.1f}%</b>", small_style),
            ],
        ]
        good_metrics_table = Table(good_metrics_data, colWidths=[80, 80, 160, 80])
        good_metrics_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#f9f9f9")),
            ("GRID", (0, 0), (-1, -1), 0.75, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(good_metrics_table)
        story.append(Spacer(1, 12))

        story.append(section_heading("Justification"))
        story.append(Spacer(1, 8))

        table_data = [[Paragraph("<b>Name</b>", normal_style), Paragraph("<b>Justification</b>", normal_style)]]
        found_rows = 0

        for item in per_intent:
            if item.get("musthave") and item.get("status") != "PRESENT":
                name = item.get("tag", "") or item.get("intent", "") or "N/A"
                status_val = item.get("status", "")
                justification = item.get("justification", "") or "No justification available."
                justification_text = f"[{status_val}] {justification}" if status_val else justification
                table_data.append([
                    Paragraph(safe_pdf_text(name), normal_style),
                    Paragraph(safe_pdf_text(justification_text), normal_style),
                ])
                found_rows += 1

        if found_rows == 0:
            table_data.append([
                Paragraph("No Gap", normal_style),
                Paragraph(
                    "All must-have requirements are PRESENT, so there is no justification gap.",
                    normal_style,
                ),
            ])

        justification_table = Table(table_data, colWidths=[160, 330])
        justification_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.75, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(justification_table)
        story.append(Spacer(1, 16))

        story.append(section_heading("Architecture Overview"))
        story.append(Spacer(1, 8))
        story.append(Paragraph(safe_pdf_text(architecture_summary), normal_style))
        story.append(Spacer(1, 12))

        footer_text = f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_style = ParagraphStyle(
            "FooterStyle",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.grey,
            alignment=2,
        )
        story.append(Paragraph(safe_pdf_text(footer_text), footer_style))

        doc.build(story)
        logging.info(f"Enhanced PDF Summary Report generated: {pdf_output_path}")
        return True

    except Exception as e:
        logging.error(f"Failed to generate PDF summary report: {e}")
        return False


# ============================================================================
# CLI
# ============================================================================

def main():
    """Command-line interface."""
    ap = argparse.ArgumentParser(
        description="Quality-aware Validation using Template JSON + Template PDF Evidence (text+diagrams)"
    )
    ap.add_argument("--template-json", required=True, help="Path to Template JSON (e.g., Template_V5.json)")
    ap.add_argument("--template-pdf", required=True, help="Path to Template PDF/DOCX with instructions")
    ap.add_argument("--nar-id", required=True)
    ap.add_argument("--application-name", required=True)
    ap.add_argument("--release-number", required=True)
    ap.add_argument("--rtype", required=True, help="indev | new | migration | rehost | normal")
    ap.add_argument("--top-k", type=int, default=6, help="Top-K evidence text chunks per tag")
    ap.add_argument("--scope-doc-id", default="", help="Optional: restrict validation to a specific evidence doc_id")
    ap.add_argument("--scope-doc-hash", default="", help="Deprecated alias for --scope-doc-id")
    ap.add_argument("--enable-mm", action="store_true", help="Enable multimodal lane (diagram checks)")
    ap.add_argument("--out", default="validation_quality_result.json", help="Output JSON path")
    ap.add_argument("--pdf-out", default="summary_report.pdf", help="Output PDF path")

    args = ap.parse_args()

    rtype = resolve_rtype(args.rtype)
    logging.info(f"[CFG] Using rtype='{rtype}', TOP-K={args.top_k}, MM={args.enable_mm}")

    # Fail fast on DB connection
    get_engine()
    validate_runtime_contract(MM_DIM if args.enable_mm else 0)

    scope_doc_id = args.scope_doc_id.strip() or args.scope_doc_hash.strip() or None
    if args.scope_doc_hash.strip() and not args.scope_doc_id.strip():
        logging.warning("--scope-doc-hash is deprecated; use --scope-doc-id instead.")

    if scope_doc_id:
        logging.info(f"Using explicit doc_id scope: {scope_doc_id}")
    else:
        logging.info(
            "No doc_id scope provided; validating across matching evidence for the selected NAR/release/rtype."
        )

    result = run_validation(
        template_json_path=args.template_json,
        template_pdf_path=args.template_pdf,
        nar_id=args.nar_id,
        application_name=args.application_name,
        release_number=args.release_number,
        rtype=rtype,
        top_k=args.top_k,
        scope_doc_id=scope_doc_id,
        enable_mm=args.enable_mm,
    )

    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    logging.info(f"Done. JSON: {args.out}")

    pdf_output_path = args.pdf_out if args.pdf_out.lower().endswith(".pdf") else f"{args.pdf_out}.pdf"
    pdf_ok = generate_summary_pdf(result, pdf_output_path)

    if pdf_ok:
        logging.info(f"Done. PDF: {pdf_output_path}")
    else:
        raise RuntimeError(f"PDF generation failed: {pdf_output_path}")


if __name__ == "__main__":
    main()
