#!/usr/bin/env python3
"""
Evidence Ingestion (TEXT + IMAGES) with folder-derived source_id/category.

Tables (updated DDL):
  vs_evidence(id, nar_id, application_name, release_number, rtype, embedding(768), 
              chunk, doc_version, doc_hash, doc_id, doc_uri, page_num, chunk_index, 
              chunk_hash, metadata, created_at)
  vs_mm_evidence_assets(id, nar_id, release_number, rtype, embedding(1408), caption, 
                        doc_uri, doc_version, doc_hash, doc_id, page_num, chunk_index, 
                        metadata, created_at)

Key guarantees:
  - NO page snapshots (no synthetic PNGs).
  - Robust source/category derivation from folder layout.
  - Safe purge by nar_id/release_number OR purge-all (and EXIT immediately).
    - Stable doc_hash/doc_id per file and stable chunk_hash per stored row.
    - Idempotent inserts via chunk_hash and ON CONFLICT (chunk_hash) DO NOTHING.

Examples:
  # Purge only (no ingestion)
  python3 evidence_loader.py --purge-all
  python3 evidence_loader.py --purge-scope --nar-id 123 --release-number R1
  
  # Text only
  python3 evidence_loader.py --root ./evidence --nar-id 123 --application-name Deploy-IQ --release-number R1.0
"""

import os
import json
import uuid
import hashlib
import base64
import argparse
import logging
import atexit
import inspect
import re
import sys
from dataclasses import dataclass
from io import StringIO
from typing import List, Dict, Any, Optional, Tuple, Iterator
import csv

import numpy as np

# Database
import sqlalchemy
from sqlalchemy import text
from google.cloud.sql.connector import Connector, IPTypes

# Vertex AI
from google import genai
from google.genai import types as genai_types
import vertexai

# Auth for REST predict (multimodal)
import google.auth
from google.auth.transport.requests import Request as GARequest
import requests

# Loaders
import docx
import PyPDF2
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import fitz  # PyMuPDF

# Chunking & similarity merge
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sklearn.metrics.pairwise import cosine_similarity


# ============================================================================
# CONFIG
# ============================================================================

PROJECT_ID = os.getenv("VERTEX_PROJECT_ID")
REGION = os.getenv("VERTEX_REGION", "europe-west3")
EMBED_MODEL = os.getenv("VERTEX_EMBED_MODEL", "text-embedding-005")
MM_EMBED_MODEL = os.getenv("VERTEX_MM_EMBED_MODEL", "multimodalembedding@001")
MM_DIM = int(os.getenv("VERTEX_MM_DIM", "1408"))

INSTANCE_CONNECTION_NAME = os.getenv("INSTANCE_CONNECTION_NAME", "")
PGUSER = os.getenv("PGUSER", os.getenv("PG_USER", "master"))
PGPASSWORD = os.getenv("PGPASSWORD", os.getenv("PG_PASSWORD", ""))
PGDATABASE = os.getenv("PGDATABASE", os.getenv("PG_DB", "database"))

TABLE_EVIDENCE = "vs_evidence"
TABLE_MM_EVIDENCE = "vs_mm_evidence_assets"
TABLE_EVIDENCE_TEXT_SOURCES = "vs_evidence_text_sources"

# Default rtype values (can be overridden via CLI)
DEFAULT_RTYPE_TEXT = os.getenv("EVIDENCE_RTYPE_TEXT", "evidence-text")
DEFAULT_RTYPE_IMAGE = os.getenv("EVIDENCE_RTYPE_IMAGE", "evidence-image")

# Embedding dimensionality
EMBED_DIM_MAP: Dict[str, int] = {
    "text-embedding-005": 768,
    "gemini-embedding-001": 3072,
}
EMBED_DIM: Optional[int] = EMBED_DIM_MAP.get(EMBED_MODEL)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s"
)

if not PROJECT_ID:
    logging.warning("VERTEX_PROJECT_ID is not set.")

vertexai.init(project=PROJECT_ID, location=REGION)
_genai = genai.Client(vertexai=True, project=PROJECT_ID, location=REGION)

ALLOWED_RTYPES = {
    "indev": "New (indev) releases",
    "new": "GCP New releases",
    "migration": "GCP Migration releases",
    "rehost": "GCP Re-Host releases",
    "normal": "GCP Normal releases",
}


# ============================================================================
# DB Engine
# ============================================================================

_connector: Optional[Connector] = None
_engine: Optional[sqlalchemy.Engine] = None


def build_engine() -> sqlalchemy.Engine:
    """Build SQLAlchemy engine via Cloud SQL Connector."""
    global _connector
    _connector = Connector()

    def getconn():
        # Private Service Connect (consistent with your validator)
        return _connector.connect(
            INSTANCE_CONNECTION_NAME,
            "pg8000",
            user=PGUSER,
            password=PGPASSWORD,
            db=PGDATABASE,
            ip_type=IPTypes.PSC,
        )

    return sqlalchemy.create_engine(
        "postgresql+pg8000://",
        creator=getconn,
        pool_pre_ping=True,
        pool_recycle=300,
    )


def get_engine() -> sqlalchemy.Engine:
    """Get or create the shared SQLAlchemy engine."""
    global _engine
    if _engine is None:
        if not INSTANCE_CONNECTION_NAME:
            raise RuntimeError("INSTANCE_CONNECTION_NAME is required.")
        _engine = build_engine()
        logging.info("SQLAlchemy engine created via Cloud SQL Connector (pg8000).")
    return _engine


def _get_column_format_type(table_name: str, column_name: str) -> Optional[str]:
    """Return PostgreSQL's formatted column type, e.g. vector(1408)."""
    engine = get_engine()
    sql = text("""
        SELECT format_type(a.atttypid, a.atttypmod) AS format_type
        FROM pg_attribute a
        JOIN pg_class c ON c.oid = a.attrelid
        JOIN pg_namespace n ON n.oid = c.relnamespace
        WHERE c.relname = :table_name
          AND a.attname = :column_name
          AND a.attnum > 0
          AND NOT a.attisdropped
        ORDER BY CASE WHEN n.nspname = current_schema() THEN 0 ELSE 1 END, n.nspname
        LIMIT 1
    """)
    with engine.connect() as conn:
        row = conn.execute(sql, {"table_name": table_name, "column_name": column_name}).first()
    return row[0] if row and row[0] else None


def _parse_vector_dimension(format_type_name: Optional[str]) -> Optional[int]:
    """Parse a vector(N) type string into its dimension."""
    if not format_type_name:
        return None
    match = re.fullmatch(r"vector\((\d+)\)", format_type_name.strip(), flags=re.IGNORECASE)
    return int(match.group(1)) if match else None


def _has_unique_chunk_hash_index(table_name: str) -> bool:
    """Check whether a table has a single-column UNIQUE index/constraint on chunk_hash."""
    engine = get_engine()
    sql = text("""
        SELECT 1
        FROM pg_index idx
        JOIN pg_class tbl ON tbl.oid = idx.indrelid
        JOIN pg_namespace ns ON ns.oid = tbl.relnamespace
        JOIN LATERAL unnest(idx.indkey) WITH ORDINALITY AS cols(attnum, ord) ON true
        JOIN pg_attribute attr ON attr.attrelid = tbl.oid AND attr.attnum = cols.attnum
        WHERE tbl.relname = :table_name
          AND ns.nspname = current_schema()
          AND idx.indisunique
        GROUP BY idx.indexrelid
        HAVING COUNT(*) = 1 AND max(attr.attname) = 'chunk_hash'
        LIMIT 1
    """)
    with engine.connect() as conn:
        return conn.execute(sql, {"table_name": table_name}).first() is not None


def _quote_ident(identifier: str) -> str:
    return '"' + identifier.replace('"', '""') + '"'


def _find_single_column_unique_constraints(table_name: str, column_name: str) -> List[str]:
    """Return single-column UNIQUE constraints for a specific column in the current schema."""
    engine = get_engine()
    sql = text("""
        SELECT con.conname
        FROM pg_constraint con
        JOIN pg_class tbl ON tbl.oid = con.conrelid
        JOIN pg_namespace ns ON ns.oid = tbl.relnamespace
        JOIN pg_attribute attr ON attr.attrelid = tbl.oid AND attr.attnum = con.conkey[1]
        WHERE tbl.relname = :table_name
          AND ns.nspname = current_schema()
          AND con.contype = 'u'
          AND array_length(con.conkey, 1) = 1
          AND attr.attname = :column_name
        ORDER BY con.conname
    """)
    with engine.connect() as conn:
        rows = conn.execute(sql, {"table_name": table_name, "column_name": column_name}).fetchall()
    return [row[0] for row in rows if row and row[0]]


def _find_single_column_unique_indexes(table_name: str, column_name: str) -> List[str]:
    """Return standalone single-column UNIQUE indexes for a specific column in the current schema."""
    engine = get_engine()
    sql = text("""
        SELECT idx_cls.relname
        FROM pg_index idx
        JOIN pg_class tbl ON tbl.oid = idx.indrelid
        JOIN pg_namespace ns ON ns.oid = tbl.relnamespace
        JOIN pg_class idx_cls ON idx_cls.oid = idx.indexrelid
        JOIN LATERAL unnest(idx.indkey) WITH ORDINALITY AS cols(attnum, ord) ON true
        JOIN pg_attribute attr ON attr.attrelid = tbl.oid AND attr.attnum = cols.attnum
        LEFT JOIN pg_constraint con ON con.conindid = idx.indexrelid
        WHERE tbl.relname = :table_name
          AND ns.nspname = current_schema()
          AND idx.indisunique
          AND con.oid IS NULL
        GROUP BY idx.indexrelid, idx_cls.relname
        HAVING COUNT(*) = 1 AND max(attr.attname) = :column_name
        ORDER BY idx_cls.relname
    """)
    with engine.connect() as conn:
        rows = conn.execute(sql, {"table_name": table_name, "column_name": column_name}).fetchall()
    return [row[0] for row in rows if row and row[0]]


def _drop_legacy_single_column_uniques(table_name: str, column_name: str) -> None:
    """Drop obsolete single-column UNIQUE constraints/indexes that conflict with row-level ingestion."""
    constraint_names = _find_single_column_unique_constraints(table_name, column_name)
    index_names = _find_single_column_unique_indexes(table_name, column_name)
    if not constraint_names and not index_names:
        return

    table_ident = _quote_ident(table_name)
    try:
        engine = get_engine()
        with engine.begin() as conn:
            for constraint_name in constraint_names:
                conn.execute(text(
                    f"ALTER TABLE {table_ident} DROP CONSTRAINT IF EXISTS {_quote_ident(constraint_name)}"
                ))
            for index_name in index_names:
                conn.execute(text(f"DROP INDEX IF EXISTS {_quote_ident(index_name)}"))
    except Exception as exc:
        legacy_names = ", ".join(constraint_names + index_names)
        raise RuntimeError(
            f"Failed to drop legacy UNIQUE constraint/index on {table_name}({column_name}) [{legacy_names}]: {exc}"
        ) from exc

    remaining_constraints = _find_single_column_unique_constraints(table_name, column_name)
    remaining_indexes = _find_single_column_unique_indexes(table_name, column_name)
    if remaining_constraints or remaining_indexes:
        legacy_names = ", ".join(remaining_constraints + remaining_indexes)
        raise RuntimeError(
            f"Legacy UNIQUE constraint/index still exists on {table_name}({column_name}) after drop attempt: {legacy_names}"
        )

    logging.info(
        f"[SCHEMA] Dropped legacy UNIQUE constraint/index on {table_name}({column_name}): "
        f"{', '.join(constraint_names + index_names)}"
    )


def _find_duplicate_chunk_hashes(
    table_name: str,
    limit: int = 5,
) -> List[Tuple[str, int]]:
    """Return a small sample of duplicate non-null chunk_hash values."""
    engine = get_engine()
    sql = text(f"""
        SELECT chunk_hash, COUNT(*) AS dup_count
        FROM {table_name}
        WHERE chunk_hash IS NOT NULL
        GROUP BY chunk_hash
        HAVING COUNT(*) > 1
        ORDER BY dup_count DESC, chunk_hash
        LIMIT :limit
    """)
    with engine.connect() as conn:
        rows = conn.execute(sql, {"limit": limit}).fetchall()
    return [(row[0], row[1]) for row in rows if row and row[0]]


def _ensure_unique_chunk_hash_index(table_name: str) -> None:
    """Create the required unique index on chunk_hash when it is missing."""
    if _has_unique_chunk_hash_index(table_name):
        return

    duplicates = _find_duplicate_chunk_hashes(table_name)
    if duplicates:
        sample = ", ".join(f"{chunk_hash[:12]}... (x{count})" for chunk_hash, count in duplicates)
        raise RuntimeError(
            f"Cannot create UNIQUE chunk_hash index for {table_name} because duplicate chunk_hash values already exist. "
            f"Sample duplicates: {sample}. Clean duplicates or purge the table/scope, then retry."
        )

    index_name = f"ux_{table_name}_chunk_hash"
    engine = get_engine()
    try:
        with engine.begin() as conn:
            conn.execute(text(f"CREATE UNIQUE INDEX IF NOT EXISTS {index_name} ON {table_name} (chunk_hash)"))
    except Exception as exc:
        raise RuntimeError(
            f"Failed to create required UNIQUE chunk_hash index '{index_name}' on {table_name}: {exc}"
        ) from exc

    if not _has_unique_chunk_hash_index(table_name):
        raise RuntimeError(
            f"Expected UNIQUE chunk_hash index on {table_name} after creation attempt, but none was detected."
        )

    logging.info(f"[SCHEMA] Ensured UNIQUE chunk_hash index on {table_name}(chunk_hash).")


def ensure_support_tables() -> None:
    """Create supporting tables used for exact text provenance when missing."""
    engine = get_engine()
    with engine.begin() as conn:
        conn.execute(text(f"""
            CREATE TABLE IF NOT EXISTS {TABLE_EVIDENCE_TEXT_SOURCES} (
                id uuid PRIMARY KEY,
                nar_id text NOT NULL,
                release_number text NOT NULL,
                doc_hash text NOT NULL,
                doc_id text NOT NULL,
                doc_uri text NOT NULL,
                page_num integer NULL,
                text_scope text NOT NULL,
                extractor_name text NOT NULL,
                parser_version text NULL,
                canonical_text text NOT NULL,
                canonical_text_hash text NOT NULL,
                source_hash text NOT NULL,
                metadata jsonb NOT NULL DEFAULT '{{}}'::jsonb,
                created_at timestamptz NOT NULL DEFAULT now()
            )
        """))
        conn.execute(text(
            f"CREATE UNIQUE INDEX IF NOT EXISTS ux_{TABLE_EVIDENCE_TEXT_SOURCES}_source_hash "
            f"ON {TABLE_EVIDENCE_TEXT_SOURCES} (source_hash)"
        ))
        conn.execute(text(
            f"CREATE INDEX IF NOT EXISTS ix_{TABLE_EVIDENCE_TEXT_SOURCES}_scope "
            f"ON {TABLE_EVIDENCE_TEXT_SOURCES} (nar_id, release_number, doc_hash, page_num)"
        ))


def validate_runtime_contract(mm_dim: int) -> None:
    """Fail fast when DB schema and runtime config cannot guarantee idempotent ingestion."""
    if EMBED_DIM is not None:
        text_embedding_type = _get_column_format_type(TABLE_EVIDENCE, "embedding")
        text_embedding_dim = _parse_vector_dimension(text_embedding_type)
        if text_embedding_dim != EMBED_DIM:
            raise RuntimeError(
                f"{TABLE_EVIDENCE}.embedding expects {text_embedding_type or 'unknown'}, "
                f"but {EMBED_MODEL} is configured for dimension {EMBED_DIM}."
            )

    mm_embedding_type = _get_column_format_type(TABLE_MM_EVIDENCE, "embedding")
    mm_embedding_dim = _parse_vector_dimension(mm_embedding_type)
    if mm_embedding_dim != mm_dim:
        raise RuntimeError(
            f"{TABLE_MM_EVIDENCE}.embedding expects {mm_embedding_type or 'unknown'}, "
            f"but VERTEX_MM_DIM is configured for dimension {mm_dim}."
        )

    for table_name in (TABLE_EVIDENCE, TABLE_MM_EVIDENCE):
        for column_name in ("doc_hash", "doc_id"):
            _drop_legacy_single_column_uniques(table_name, column_name)
        _ensure_unique_chunk_hash_index(table_name)
    ensure_support_tables()


def _shutdown_connector():
    """Cleanup: close Cloud SQL connector."""
    global _connector, _engine
    try:
        if _engine is not None:
            _engine.dispose()
    finally:
        if _connector is not None:
            _connector.close()
        logging.info("Cloud SQL Connector closed.")


atexit.register(_shutdown_connector)


# ============================================================================
# Helpers
# ============================================================================


@dataclass(frozen=True)
class ExtractedTextSource:
    text_scope: str
    page_num: Optional[int]
    text: str
    extractor_name: str
    parser_version: Optional[str]


@dataclass(frozen=True)
class ChunkRecord:
    text: str
    start_offset: int
    end_offset: int


@dataclass(frozen=True)
class CsvChunkSource:
    chunk_index: int
    row_start: int
    row_end: int
    text: str


@dataclass(frozen=True)
class ExtractedImageAsset:
    page_num: Optional[int]
    chunk_index: int
    img_bytes: bytes
    caption: str
    asset_label: str
    nearby_heading: Optional[str] = None
    preceding_text: Optional[str] = None
    page_text_excerpt: Optional[str] = None


def _version_string(value: Any) -> Optional[str]:
    if value in (None, ""):
        return None
    if isinstance(value, tuple):
        return ".".join(str(part) for part in value if part not in (None, ""))
    return str(value)


def _module_version(module: Any, *attrs: str) -> Optional[str]:
    for attr in attrs:
        value = _version_string(getattr(module, attr, None))
        if value:
            return value
    return None


PYPDF2_VERSION = _module_version(PyPDF2, "__version__", "version")
PYMUPDF_VERSION = _module_version(fitz, "VersionBind", "__version__")
DOCX_VERSION = _module_version(docx, "__version__", "VERSION")
PYTHON_RUNTIME_VERSION = sys.version.split()[0]

def _sha_bytes(b: bytes) -> str:
    """Compute SHA-256 hash of bytes."""
    return hashlib.sha256(b).hexdigest()


def _sha_text(s: str) -> str:
    """Compute SHA-256 hash of text."""
    return hashlib.sha256(s.encode("utf-8")).hexdigest()


def _sha_parts(*parts: Any) -> str:
    """Compute SHA-256 over typed, length-delimited parts."""
    h = hashlib.sha256()
    for part in parts:
        if part is None:
            kind = b"N"
            payload = b""
        elif isinstance(part, bytes):
            kind = b"B"
            payload = part
        else:
            kind = b"T"
            payload = str(part).encode("utf-8")
        h.update(kind)
        h.update(len(payload).to_bytes(8, "big"))
        h.update(payload)
    return h.hexdigest()


def _vec_literal(vec: List[float]) -> str:
    """Format vector as PostgreSQL pgvector literal."""
    return "[" + ",".join(f"{x:.6f}" for x in vec) + "]"


def _normalize_context_text(text_in: str) -> str:
    """Normalize extracted layout text for compact metadata storage."""
    return re.sub(r"\s+", " ", (text_in or "").replace("\x00", " ")).strip()


def _truncate_context_text(text_in: str, max_chars: int) -> str:
    """Trim context text to a stable maximum length."""
    normalized = _normalize_context_text(text_in)
    if len(normalized) <= max_chars:
        return normalized
    return normalized[: max_chars - 3].rstrip() + "..."


def _extract_pdf_text_block_text(block: Dict[str, Any]) -> str:
    """Flatten a PyMuPDF text block into readable text."""
    lines: List[str] = []
    for line in block.get("lines") or []:
        spans: List[str] = []
        for span in line.get("spans") or []:
            span_text = _normalize_context_text(span.get("text", ""))
            if span_text:
                spans.append(span_text)
        if spans:
            lines.append(" ".join(spans))
    if lines:
        return "\n".join(lines)
    return _normalize_context_text(str(block.get("text") or ""))


def _looks_like_heading_candidate(text_in: str) -> bool:
    """Heuristically detect section headings near diagrams."""
    raw_lines = [
        _normalize_context_text(line)
        for line in re.split(r"\r?\n", text_in or "")
    ]
    raw_lines = [line for line in raw_lines if line]
    if not raw_lines or len(raw_lines) > 3:
        return False
    normalized = " ".join(raw_lines).strip()
    if len(normalized) < 4 or len(normalized) > 180:
        return False
    lower = normalized.lower()
    if re.match(r"^[A-Za-z]?\d+(?:[.\s]+\d+)*(?:\.[a-z])?\s+\S", normalized):
        return True
    if any(token in lower for token in ("diagram", "architecture", "model", "flow", "component", "interface")):
        return True
    if normalized.endswith(":"):
        return True
    word_count = len(normalized.split())
    if word_count <= 12 and not normalized.endswith("."):
        alpha_chars = sum(1 for char in normalized if char.isalpha())
        return alpha_chars >= max(3, int(len(normalized) * 0.55))
    return False


def _select_nearest_heading(block_texts: List[str]) -> Optional[str]:
    """Choose the nearest likely heading from preceding text blocks."""
    for text in reversed(block_texts[-6:]):
        if _looks_like_heading_candidate(text):
            return _truncate_context_text(text, 180)
    return None


def _build_preceding_text_context(
    block_texts: List[str],
    nearby_heading: Optional[str],
) -> Optional[str]:
    """Build a short nearby-text context window for an image asset."""
    heading_norm = _normalize_context_text(nearby_heading or "")
    recent_blocks: List[str] = []
    for text in block_texts[-3:]:
        normalized = _normalize_context_text(text)
        if not normalized or normalized == heading_norm:
            continue
        recent_blocks.append(normalized)
    if not recent_blocks:
        return nearby_heading
    return _truncate_context_text(" | ".join(recent_blocks), 220)


def _build_image_context_caption(
    asset_label: str,
    nearby_heading: Optional[str],
    preceding_text: Optional[str],
) -> str:
    """Build a caption that carries nearby section context into the MM embedding."""
    caption_parts: List[str] = []
    if nearby_heading:
        caption_parts.append(_normalize_context_text(nearby_heading))
    normalized_preceding = _normalize_context_text(preceding_text or "")
    if normalized_preceding and normalized_preceding not in caption_parts:
        caption_parts.append(normalized_preceding)
    if not caption_parts:
        return asset_label
    return _truncate_context_text(" | ".join(caption_parts), 220)


def _render_pdf_structured_page_text(page: Any) -> str:
    """Render page text from ordered PyMuPDF blocks while preserving layout boundaries."""
    try:
        blocks = (page.get_text("dict", sort=True) or {}).get("blocks") or []
    except Exception:
        return ""

    rendered_blocks: List[str] = []
    for block in blocks:
        if block.get("type") != 0:
            continue
        block_text = _extract_pdf_text_block_text(block)
        if not block_text:
            continue
        lines = [
            _normalize_context_text(line)
            for line in re.split(r"\r?\n", block_text)
        ]
        lines = [line for line in lines if line]
        if not lines:
            continue
        rendered_blocks.append("\n".join(lines))

    normalized_page_text = _normalize_context_text("\n\n".join(rendered_blocks))
    for table_text in _extract_pdf_table_texts(page):
        normalized_table_text = _normalize_context_text(table_text)
        if not normalized_table_text or normalized_table_text in normalized_page_text:
            continue
        rendered_blocks.append("TABLE\n" + table_text)
        normalized_page_text = _normalize_context_text("\n\n".join(rendered_blocks))
    return "\n\n".join(rendered_blocks).strip()


def _normalized_text_length(text_in: str) -> int:
    """Measure extracted text content after normalization for extractor selection."""
    return len(re.sub(r"\s+", "", (text_in or "").strip()))


def _choose_pdf_page_text(
    pypdf_text: str,
    fitz_structured_text: str,
    fitz_plain_text: str,
) -> Tuple[str, str, Optional[str]]:
    """Choose the best page text candidate while preferring structured layout output."""
    structured_len = _normalized_text_length(fitz_structured_text)
    pypdf_len = _normalized_text_length(pypdf_text)
    fitz_plain_len = _normalized_text_length(fitz_plain_text)

    if structured_len and (pypdf_len == 0 or structured_len >= max(80, int(pypdf_len * 0.6))):
        return fitz_structured_text, "pymupdf-layout", PYMUPDF_VERSION
    if pypdf_len >= fitz_plain_len and pypdf_len:
        return pypdf_text, "pypdf2", PYPDF2_VERSION
    if fitz_plain_len:
        return fitz_plain_text, "pymupdf", PYMUPDF_VERSION
    return "", "", None


def _extract_pdf_table_texts(page: Any) -> List[str]:
    """Extract table rows when PyMuPDF table detection is available."""
    if not hasattr(page, "find_tables"):
        return []
    try:
        table_result = page.find_tables()
    except Exception:
        return []

    tables = getattr(table_result, "tables", table_result) or []
    extracted_tables: List[str] = []
    for table in tables:
        try:
            rows = table.extract() or []
        except Exception:
            continue
        rendered_rows: List[str] = []
        for row in rows:
            if not isinstance(row, (list, tuple)):
                continue
            cells = [_normalize_context_text(cell or "") for cell in row]
            if any(cells):
                rendered_rows.append(" | ".join(cells))
        if rendered_rows:
            extracted_tables.append("\n".join(rendered_rows))
    return extracted_tables


def resolve_rtype(user_val: str) -> str:
    """Maps user-provided shorthand to canonical rtype string.
    
    Accepts: indev, new, migration, rehost, normal
    """
    key = user_val.strip().lower()
    if key in ALLOWED_RTYPES:
        return ALLOWED_RTYPES[key]
    # Allow full exact match too
    for canonical in ALLOWED_RTYPES.values():
        if key == canonical.lower():
            return canonical
    raise ValueError(
        f"Invalid rtype '{user_val}'. Allowed: indev, new, migration, rehost, normal"
    )


# Allowed evidence file types (including CSV)
ALLOWED_EXTS = (".pdf", ".docx", ".txt", ".csv")


def _resolve_file_paths(root: str, paths: List[str]) -> List[str]:
    """Resolves a list of file paths to absolute paths.
    
    If a path is absolute, use as-is.
    If a path is relative, resolve relative to root.
    Filters to allowed extensions.
    Ensures files actually exist.
    """
    out: List[str] = []
    for p in paths:
        p = p.strip()
        if not p:
            continue
        abs_path = p if os.path.isabs(p) else os.path.abspath(os.path.join(root, p))
        if not os.path.isfile(abs_path):
            logging.warning(f"[ONLY-FILE] Skipping non-existent file: {abs_path}")
            continue
        if not abs_path.lower().endswith(ALLOWED_EXTS):
            logging.warning(f"[ONLY-FILE] skipping unsupported file type: {abs_path}")
            continue
        out.append(abs_path)
    return out


# ============================================================================
# Text Embeddings
# ============================================================================

def _extract_values(resp) -> List[float]:
    """Extract embedding vector from various response formats."""
    if hasattr(resp, "embedding") and hasattr(resp.embedding, "values"):
        return resp.embedding.values
    if hasattr(resp, "values"):
        return resp.values
    if hasattr(resp, "embeddings") and resp.embeddings:
        for e in resp.embeddings:
            if hasattr(e, "values"):
                return e.values
            if isinstance(e, dict) and "values" in e:
                return e["values"]
    if isinstance(resp, dict):
        if "embedding" in resp and isinstance(resp["embedding"], dict) and "values" in resp["embedding"]:
            return resp["embedding"]["values"]
        if "values" in resp:
            return resp["values"]
    raise RuntimeError(f"Unexpected embed response shape: {type(resp)}")


def embed_texts(
    texts: List[str],
    task_type: str = "RETRIEVAL_DOCUMENT",
    out_dim: Optional[int] = EMBED_DIM,
) -> List[List[float]]:
    """Embed multiple text chunks."""
    cfg = genai_types.EmbedContentConfig(
        task_type=task_type,
        output_dimensionality=out_dim,
    )
    out: List[List[float]] = []
    for t in texts:
        r = _genai.models.embed_content(model=EMBED_MODEL, contents=t, config=cfg)
        out.append(_extract_values(r))
    return out


# ============================================================================
# Multimodal Embeddings (REST)
# ============================================================================

def get_access_token() -> str:
    """Get Google Cloud access token for REST API calls."""
    creds, _ = google.auth.default(
        scopes=["https://www.googleapis.com/auth/cloud-platform"]
    )
    if not creds.valid:
        creds.refresh(GARequest())
    return creds.token


def mm_embed_image_with_caption(
    img_bytes: bytes,
    caption: str,
    dim: int = MM_DIM,
) -> List[float]:
    """Embed image with optional caption using multimodal embedding model."""
    token = get_access_token()
    endpoint = (
        f"https://{REGION}-aiplatform.googleapis.com/v1/projects/{PROJECT_ID}/locations/{REGION}"
        f"/publishers/google/models/{MM_EMBED_MODEL}:predict"
    )
    b64 = base64.b64encode(img_bytes).decode("utf-8")
    instance = {
        "image": {
            "bytesBase64Encoded": b64,
            "mimeType": "image/png",
        },
        "parameters": {"dimension": dim},
    }
    if caption:
        instance["text"] = caption[:200]
    payload = {"instances": [instance]}
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }
    resp = requests.post(endpoint, headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    data = resp.json()
    preds = data.get("predictions") or []
    if not preds:
        raise RuntimeError(f"Vertex MM: missing predictions. keys={list(data.keys())}")
    pred = preds[0]
    vec = pred.get("imageEmbedding") or pred.get("textEmbedding")
    if vec is None:
        raise RuntimeError(f"Vertex MM: missing embedding. pred keys={list(pred.keys())}")
    if dim and len(vec) != dim:
        raise RuntimeError(f"Vertex MM: expected dim={dim}, got len={len(vec)}")
    return vec


# ============================================================================
# Chunking & Semantic Merge (parity)
# ============================================================================


def _splitter_supports_start_index() -> bool:
    try:
        return "add_start_index" in inspect.signature(
            RecursiveCharacterTextSplitter.__init__
        ).parameters
    except (TypeError, ValueError):
        return False

def get_splitter(
    chunk_size: int = 1000,
    chunk_overlap: int = 100,
    add_start_index: bool = False,
) -> RecursiveCharacterTextSplitter:
    """Create a text splitter."""
    kwargs: Dict[str, Any] = {
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "separators": ["\n\n", "\n", " ", ""],
    }
    if add_start_index and _splitter_supports_start_index():
        kwargs["add_start_index"] = True
    return RecursiveCharacterTextSplitter(**kwargs)


def _map_chunks_to_offsets(
    text_in: str,
    raw_chunks: List[str],
    chunk_overlap: int,
) -> List[ChunkRecord]:
    """Map splitter output back to deterministic offsets in the canonical source text."""
    records: List[ChunkRecord] = []
    search_start = 0
    for chunk in raw_chunks:
        if not chunk:
            continue
        start = text_in.find(chunk, search_start)
        if start == -1 and records:
            fallback_start = max(0, records[-1].end_offset - chunk_overlap - len(chunk))
            start = text_in.find(chunk, fallback_start)
        if start == -1:
            start = text_in.find(chunk)
        if start == -1:
            raise RuntimeError("Failed to map a chunk back to canonical source text.")
        end = start + len(chunk)
        records.append(ChunkRecord(text=chunk, start_offset=start, end_offset=end))
        search_start = max(0, end - chunk_overlap)
    return records


def split_text_records(
    text_in: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 100,
) -> List[ChunkRecord]:
    """Split text into deterministic chunk records with offsets."""
    if not text_in or not text_in.strip():
        return []
    if _splitter_supports_start_index():
        docs = get_splitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            add_start_index=True,
        ).create_documents([text_in])
        records: List[ChunkRecord] = []
        for doc in docs:
            chunk_text = doc.page_content
            start_offset = int(doc.metadata.get("start_index", 0))
            records.append(
                ChunkRecord(
                    text=chunk_text,
                    start_offset=start_offset,
                    end_offset=start_offset + len(chunk_text),
                )
            )
        return records
    raw_chunks = get_splitter(chunk_size, chunk_overlap).split_text(text_in)
    return _map_chunks_to_offsets(text_in, raw_chunks, chunk_overlap)


def merge_similar_chunk_records(
    text_in: str,
    chunk_records: List[ChunkRecord],
    threshold: float = 0.85,
) -> List[ChunkRecord]:
    """Merge adjacent chunks while preserving exact source slices and offsets."""
    if len(chunk_records) < 2:
        return chunk_records
    embeddings = embed_texts(
        [record.text for record in chunk_records],
        task_type="RETRIEVAL_DOCUMENT",
        out_dim=EMBED_DIM,
    )
    merged: List[ChunkRecord] = []
    current_record = chunk_records[0]
    cur_emb = np.array(embeddings[0], dtype=np.float32)
    for next_record, next_emb_values in zip(chunk_records[1:], embeddings[1:]):
        next_emb = np.array(next_emb_values, dtype=np.float32)
        sim = cosine_similarity([cur_emb], [next_emb])[0][0]
        if sim >= threshold:
            current_record = ChunkRecord(
                text=text_in[current_record.start_offset:next_record.end_offset],
                start_offset=current_record.start_offset,
                end_offset=next_record.end_offset,
            )
            cur_emb = (cur_emb + next_emb) / 2.0
        else:
            merged.append(current_record)
            current_record = next_record
            cur_emb = next_emb
    merged.append(current_record)
    return merged


def build_chunk_records(
    text_in: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 100,
    merge_threshold: float = 0.85,
    semantic_merge: bool = False,
) -> List[ChunkRecord]:
    """Build chunk records with deterministic offsets and optional semantic merge."""
    records = split_text_records(
        text_in,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    if semantic_merge:
        return merge_similar_chunk_records(text_in, records, threshold=merge_threshold)
    return records


def _serialize_csv_chunk(header: Optional[List[str]], rows: List[List[str]]) -> str:
    """Serialize CSV rows canonically while preserving valid CSV quoting."""
    buf = StringIO()
    writer = csv.writer(buf, lineterminator="\n")
    if header is not None:
        writer.writerow(header)
    writer.writerows(rows)
    return buf.getvalue().rstrip("\n")


def split_csv_rows(
    file_path: str,
    rows_per_chunk: int = 50,
) -> List[CsvChunkSource]:
    """Split a CSV into canonical chunk texts, keeping the header in each chunk."""
    chunks: List[CsvChunkSource] = []
    try:
        with open(file_path, newline="", encoding="utf-8") as csvfile:
            reader = csv.reader(csvfile)
            header = next(reader, None)
            current: List[List[str]] = []
            chunk_start_row = 1
            for row_num, row in enumerate(reader, start=1):
                current.append(row)
                if row_num % rows_per_chunk == 0:
                    chunks.append(
                        CsvChunkSource(
                            chunk_index=len(chunks),
                            row_start=chunk_start_row,
                            row_end=row_num,
                            text=_serialize_csv_chunk(header, current),
                        )
                    )
                    current = []
                    chunk_start_row = row_num + 1
            if current:
                chunks.append(
                    CsvChunkSource(
                        chunk_index=len(chunks),
                        row_start=chunk_start_row,
                        row_end=chunk_start_row + len(current) - 1,
                        text=_serialize_csv_chunk(header, current),
                    )
                )
    except Exception as e:
        logging.error(f"[CSV] Failed to chunk '{file_path}': {e}")
    return chunks


def parse_csv_image_rows(
    file_path: str,
    image_columns: List[str],
    caption_column: Optional[str] = None,
) -> List[Tuple[int, int, bytes, str]]:
    """Extract images from CSV rows.
    
    Returns (row_num, img_index_in_row, img_bytes, caption).
    Supports base64-encoded PNG/JPEG in specified columns.
    """
    out: List[Tuple[int, int, bytes, str]] = []
    if not image_columns:
        return out
    try:
        with open(file_path, newline="", encoding="utf-8") as csvfile:
            reader = csv.DictReader(csvfile)
            for row_num, row in enumerate(reader, start=1):
                if not row:
                    continue
                cap = (row.get(caption_column) or "").strip() if caption_column else ""
                for idx, col in enumerate(image_columns, start=1):
                    raw = (row.get(col) or "").strip()
                    if not raw:
                        continue
                    if raw.startswith("data:image"):
                        try:
                            raw = raw.split(",", 1)[-1]
                        except Exception:
                            pass
                    try:
                        img_bytes = base64.b64decode(raw, validate=True)
                    except Exception:
                        logging.warning(
                            f"[CSV] [IMG] Row {row_num} col '{col}' not valid base64; skipping."
                        )
                        continue
                    caption = cap or f"row{row_num}_col{col}"
                    out.append((row_num, idx, img_bytes, caption))
    except Exception as e:
        logging.error(f"[CSV] [IMG] Failed to parse images in '{file_path}': {e}")
    return out


# ============================================================================
# File Walking & Identifiers
# ============================================================================

def derive_source_and_category(
    root: str,
    file_path: str,
    strict_root: bool = False,
) -> Tuple[str, Optional[str], str]:
    """Source id = first folder under root.
    Category = second folder under root (if exists).
    
    Enforces that files MUST NOT be directly under root.
    """
    rel = os.path.relpath(file_path, root).strip(os.sep)
    parts = rel.split(os.sep)
    
    # Must be at least ROOT/<source_id>/<file>
    if len(parts) < 2:
        msg = f"[SKIP] File is directly under root (no source_id folder): {file_path}"
        if strict_root:
            raise ValueError(msg + " (use root at the parent folder of source buckets)")
        logging.warning(msg)
        return "UNKNOWN", None, rel
    
    source_id = parts[0]
    category = parts[1] if len(parts) > 2 else None
    return source_id, category, rel


def iter_guidance_files(root: str) -> Iterator[str]:
    """Walk root and yield guidance files."""
    for dirpath, _, filenames in os.walk(root):
        for fn in filenames:
            if fn.startswith("~$"):  # ignore temp files
                continue
            low = fn.lower()
            if low.endswith((".pdf", ".docx", ".txt", ".csv")):
                yield os.path.join(dirpath, fn)


def compute_doc_id(file_path: str) -> str:
    """Compute stable file identity as SHA-256 of file bytes."""
    with open(file_path, "rb") as f:
        b = f.read()
    return _sha_bytes(b)


def compute_text_chunk_hash(
    nar_id: str,
    release_number: str,
    document_hash: str,
    page_num: Optional[int],
    chunk_index: int,
    start_offset: Optional[int],
    end_offset: Optional[int],
    chunk_text: str,
) -> str:
    """Compute a stable row hash for a text chunk within an ingestion scope."""
    return _sha_parts(
        "text",
        nar_id,
        release_number,
        document_hash,
        page_num,
        chunk_index,
        start_offset,
        end_offset,
        chunk_text,
    )


def compute_source_text_hash(
    nar_id: str,
    release_number: str,
    document_hash: str,
    page_num: Optional[int],
    text_scope: str,
    canonical_text: str,
) -> str:
    """Compute a stable row hash for stored canonical extracted text."""
    return _sha_parts(
        "source-text",
        nar_id,
        release_number,
        document_hash,
        page_num,
        text_scope,
        canonical_text,
    )


def compute_image_chunk_hash(
    nar_id: str,
    release_number: str,
    document_hash: str,
    page_num: Optional[int],
    chunk_index: int,
    caption: str,
    img_bytes: bytes,
) -> str:
    """Compute a stable row hash for an image chunk within an ingestion scope."""
    return _sha_parts(
        "image",
        nar_id,
        release_number,
        document_hash,
        page_num,
        chunk_index,
        caption,
        img_bytes,
    )


def compute_doc_uri(
    root: str,
    file_path: str,
    uri_prefix: Optional[str],
) -> str:
    """Compute doc_uri."""
    rel = os.path.relpath(file_path, root).replace("\\", "/")
    if uri_prefix:
        return uri_prefix.rstrip("/") + "/" + rel
    return os.path.abspath(file_path)


# ============================================================================
# Text Extraction (page-aware for PDF)
# ============================================================================

def pdf_pages_text(pdf_path: str) -> List[ExtractedTextSource]:
    """Extract text from PDF by page."""
    pages: List[ExtractedTextSource] = []
    reader: Optional[PdfReader] = None
    fitz_doc = None
    try:
        reader = PdfReader(pdf_path)
    except Exception as exc:
        logging.debug(f"[PDF][TEXT] PyPDF2 failed to open '{pdf_path}': {exc}")
    try:
        fitz_doc = fitz.open(pdf_path)
    except Exception as exc:
        logging.debug(f"[PDF][TEXT] PyMuPDF failed to open '{pdf_path}': {exc}")

    if reader is None and fitz_doc is None:
        return pages

    page_count = max(
        len(reader.pages) if reader is not None else 0,
        len(fitz_doc) if fitz_doc is not None else 0,
    )
    for i in range(page_count):
        pypdf_text = ""
        fitz_structured_text = ""
        fitz_text = ""
        if reader is not None and i < len(reader.pages):
            try:
                pypdf_text = reader.pages[i].extract_text() or ""
            except Exception as exc:
                logging.debug(f"[PDF][TEXT] PyPDF2 failed to extract text from page {i+1} of '{pdf_path}': {exc}")
        if fitz_doc is not None and i < len(fitz_doc):
            try:
                fitz_page = fitz_doc.load_page(i)
                fitz_structured_text = _render_pdf_structured_page_text(fitz_page)
                fitz_text = fitz_page.get_text("text") or ""
            except Exception as exc:
                logging.debug(f"[PDF][TEXT] PyMuPDF failed to extract text from page {i+1} of '{pdf_path}': {exc}")

        chosen_text, extractor_name, parser_version = _choose_pdf_page_text(
            pypdf_text,
            fitz_structured_text,
            fitz_text,
        )
        if chosen_text and chosen_text.strip():
            pages.append(
                ExtractedTextSource(
                    text_scope="page",
                    page_num=i + 1,
                    text=chosen_text,
                    extractor_name=extractor_name,
                    parser_version=parser_version,
                )
            )
    if fitz_doc is not None:
        fitz_doc.close()
    return pages


def build_pdf_document_source(
    page_sources: List[ExtractedTextSource],
) -> Optional[ExtractedTextSource]:
    """Build a document-level canonical PDF text source from extracted page sources."""
    if not page_sources:
        return None
    parser_versions = sorted(
        {
            f"{source.extractor_name}:{source.parser_version or 'unknown'}"
            for source in page_sources
        }
    )
    return ExtractedTextSource(
        text_scope="document",
        page_num=None,
        text="\n\n".join(source.text for source in page_sources),
        extractor_name="pdf-page-join",
        parser_version=", ".join(parser_versions),
    )


def docx_text(docx_path: str) -> str:
    """Extract text from DOCX."""
    doc = DocxDocument(docx_path)
    paragraphs = [p.text for p in doc.paragraphs]
    table_lines: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            if any(cell for cell in cells):
                table_lines.append(" | ".join(cells))
    combined = paragraphs + table_lines
    return "\n".join(combined)


def txt_text(txt_path: str) -> str:
    """Extract text from TXT file."""
    with open(txt_path, "rb") as f:
        raw = f.read()
    try:
        return raw.decode("utf-8", errors="ignore")
    except Exception:
        return raw.decode("latin-1", errors="ignore")


# ============================================================================
# Image Extraction (NO page snapshots)
# ============================================================================

def extract_pdf_images(pdf_path: str) -> List[ExtractedImageAsset]:
    """Extract embedded PDF images with nearby heading/text context when available."""
    out: List[ExtractedImageAsset] = []
    doc = fitz.open(pdf_path)
    for pno in range(len(doc)):
        page = doc[pno]
        page_images = page.get_images(full=True) or []
        if not page_images:
            continue

        try:
            blocks = (page.get_text("dict", sort=True) or {}).get("blocks") or []
        except Exception as exc:
            logging.debug(f"[PDF][IMG] Failed to extract layout blocks from page {pno + 1} of '{pdf_path}': {exc}")
            blocks = []

        preceding_text_blocks: List[str] = []
        image_index = 0
        for block in blocks:
            if block.get("type") == 0:
                block_text = _extract_pdf_text_block_text(block)
                if block_text:
                    preceding_text_blocks.append(block_text)
                continue
            if block.get("type") != 1:
                continue

            image_index += 1
            if image_index > len(page_images):
                continue

            xref = page_images[image_index - 1][0]
            pix = fitz.Pixmap(doc, xref)
            if pix.n == 4:  # CMYK -> RGB
                pix = fitz.Pixmap(fitz.csRGB, pix)
            img_bytes = pix.tobytes("png")
            asset_label = f"page{pno + 1}_img{image_index}"
            nearby_heading = _select_nearest_heading(preceding_text_blocks)
            preceding_text = _build_preceding_text_context(
                preceding_text_blocks,
                nearby_heading,
            )
            page_text_excerpt = _truncate_context_text(
                " ".join(preceding_text_blocks[-6:]),
                320,
            ) or None
            out.append(
                ExtractedImageAsset(
                    page_num=pno + 1,
                    chunk_index=image_index,
                    img_bytes=img_bytes,
                    caption=_build_image_context_caption(
                        asset_label,
                        nearby_heading,
                        preceding_text,
                    ),
                    asset_label=asset_label,
                    nearby_heading=nearby_heading,
                    preceding_text=preceding_text,
                    page_text_excerpt=page_text_excerpt,
                )
            )

        for fallback_index in range(image_index + 1, len(page_images) + 1):
            xref = page_images[fallback_index - 1][0]
            pix = fitz.Pixmap(doc, xref)
            if pix.n == 4:  # CMYK -> RGB
                pix = fitz.Pixmap(fitz.csRGB, pix)
            img_bytes = pix.tobytes("png")
            asset_label = f"page{pno + 1}_img{fallback_index}"
            nearby_heading = _select_nearest_heading(preceding_text_blocks)
            preceding_text = _build_preceding_text_context(
                preceding_text_blocks,
                nearby_heading,
            )
            page_text_excerpt = _truncate_context_text(
                " ".join(preceding_text_blocks[-6:]),
                320,
            ) or None
            out.append(
                ExtractedImageAsset(
                    page_num=pno + 1,
                    chunk_index=fallback_index,
                    img_bytes=img_bytes,
                    caption=_build_image_context_caption(
                        asset_label,
                        nearby_heading,
                        preceding_text,
                    ),
                    asset_label=asset_label,
                    nearby_heading=nearby_heading,
                    preceding_text=preceding_text,
                    page_text_excerpt=page_text_excerpt,
                )
            )
    doc.close()
    return out


def extract_docx_images(docx_path: str) -> List[Tuple[int, bytes, str]]:
    """Returns list of: (chunk_index, img_bytes, caption).
    
    DOCX has no page numbers; page_num will be NULL.
    """
    out: List[Tuple[int, bytes, str]] = []
    doc = DocxDocument(docx_path)
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    
    idx = 0
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            idx += 1
            img_bytes = rel.target_part.blob
            caption = os.path.basename(rel.target_ref) or f"docx_img{idx}"
            out.append((idx, img_bytes, caption))
    return out


def _image_extractor_details(file_ext: str) -> Tuple[str, Optional[str]]:
    """Return extraction provenance for image ingestion."""
    if file_ext == ".pdf":
        return "pymupdf-embedded-image", PYMUPDF_VERSION
    if file_ext == ".docx":
        return "python-docx-image", DOCX_VERSION
    if file_ext == ".csv":
        return "python-csv-base64", PYTHON_RUNTIME_VERSION
    return "unknown", None


def persist_source_text(
    nar_id: str,
    release_number: str,
    doc_id: str,
    doc_uri: str,
    document_hash: str,
    source: ExtractedTextSource,
    metadata: Dict[str, Any],
) -> str:
    """Persist canonical extracted text used for exact retrieval provenance."""
    engine = get_engine()
    source_hash = compute_source_text_hash(
        nar_id,
        release_number,
        document_hash,
        source.page_num,
        source.text_scope,
        source.text,
    )
    canonical_text_hash = _sha_text(source.text)
    ins_sql = text(f"""
        INSERT INTO {TABLE_EVIDENCE_TEXT_SOURCES}
        (id, nar_id, release_number, doc_hash, doc_id, doc_uri, page_num, text_scope,
         extractor_name, parser_version, canonical_text, canonical_text_hash, source_hash, metadata)
        VALUES (:id, :nar_id, :release_number, :doc_hash, :doc_id, :doc_uri, :page_num, :text_scope,
                :extractor_name, :parser_version, :canonical_text, :canonical_text_hash, :source_hash,
                CAST(:metadata AS jsonb))
        ON CONFLICT (source_hash) DO NOTHING
    """)
    with engine.begin() as conn:
        conn.execute(ins_sql, {
            "id": str(uuid.uuid4()),
            "nar_id": nar_id,
            "release_number": release_number,
            "doc_hash": document_hash,
            "doc_id": doc_id,
            "doc_uri": doc_uri,
            "page_num": source.page_num,
            "text_scope": source.text_scope,
            "extractor_name": source.extractor_name,
            "parser_version": source.parser_version,
            "canonical_text": source.text,
            "canonical_text_hash": canonical_text_hash,
            "source_hash": source_hash,
            "metadata": json.dumps(metadata),
        })
    return source_hash


# ============================================================================
# DB Inserts
# ============================================================================

def ingest_evidence_text_file(
    root: str,
    file_path: str,
    nar_id: str,
    application_name: str,
    release_number: str,
    source_id: str,
    category: Optional[str],
    rel_path: str,
    doc_id: str,
    doc_uri: str,
    rtype_text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 100,
    merge_threshold: float = 0.85,
    semantic_merge: bool = False,
    pdf_page_aware: bool = True,
    csv_rows_per_chunk: int = 50,
) -> int:
    """Inserts into vs_evidence.
    
    Returns inserted row count.
    """
    engine = get_engine()
    inserted = 0
    document_hash = doc_id
    ins_sql = text(f"""
        INSERT INTO {TABLE_EVIDENCE}
        (id, nar_id, application_name, release_number, rtype, doc_version, doc_hash, 
         doc_id, doc_uri, page_num, embedding, chunk_index, chunk, chunk_hash, metadata)
        VALUES (:id, :nar_id, :application_name, :release_number, :rtype, :doc_version, 
                :doc_hash, :doc_id, :doc_uri, :page_num, CAST(:embedding AS vector), 
                :chunk_index, :chunk, :chunk_hash, CAST(:metadata AS jsonb))
        ON CONFLICT (chunk_hash) DO NOTHING
        RETURNING chunk_hash
    """)
    
    ext = os.path.splitext(file_path)[1].lower()
    common_md = {
        "kind": "evidence-text",
        "rtype": rtype_text,
        "nar_id": nar_id,
        "application_name": application_name,
        "release_number": release_number,
        "source_id": source_id,
        "category": category,
        "doc_id": doc_id,
        "doc_hash": document_hash,
        "doc_uri": doc_uri,
        "file_name": os.path.basename(file_path),
        "rel_path": rel_path,
        "file_ext": ext.lstrip("."),
        "embedding_model": EMBED_MODEL,
        "chunking": {
            "strategy": "recursive-character",
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
            "semantic_merge": semantic_merge,
            "merge_threshold": merge_threshold if semantic_merge else None,
        },
    }
    chunk_attempts = 0

    def _build_source_metadata(
        source: ExtractedTextSource,
        extra_md: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        md = {
            "kind": "source-text",
            "nar_id": nar_id,
            "release_number": release_number,
            "application_name": application_name,
            "source_id": source_id,
            "category": category,
            "doc_id": doc_id,
            "doc_hash": document_hash,
            "doc_uri": doc_uri,
            "file_name": os.path.basename(file_path),
            "rel_path": rel_path,
            "file_ext": ext.lstrip("."),
            "text_scope": source.text_scope,
            "page_num": source.page_num,
            "extractor_name": source.extractor_name,
            "parser_version": source.parser_version,
        }
        if extra_md:
            md.update(extra_md)
        return md

    def _insert_chunk_records(
        source: ExtractedTextSource,
        source_hash: str,
        chunk_records: List[ChunkRecord],
        chunk_index_base: int = 0,
        extra_md: Optional[Dict[str, Any]] = None,
    ) -> None:
        nonlocal inserted, chunk_attempts
        if not chunk_records:
            return
        chunk_attempts += len(chunk_records)
        chunk_texts = [record.text for record in chunk_records]
        embs = embed_texts(chunk_texts, task_type="RETRIEVAL_DOCUMENT", out_dim=EMBED_DIM)
        with engine.begin() as conn:
            for offset, (record, emb) in enumerate(zip(chunk_records, embs)):
                chunk_index = chunk_index_base + offset
                md = dict(common_md)
                md.update({
                    "page_num": source.page_num,
                    "chunk_index": chunk_index,
                    "start_offset": record.start_offset,
                    "end_offset": record.end_offset,
                    "text_scope": source.text_scope,
                    "source_text_hash": source_hash,
                    "extractor_name": source.extractor_name,
                    "parser_version": source.parser_version,
                })
                if extra_md:
                    md.update(extra_md)
                chunk_hash = compute_text_chunk_hash(
                    nar_id,
                    release_number,
                    document_hash,
                    source.page_num,
                    chunk_index,
                    record.start_offset,
                    record.end_offset,
                    record.text,
                )
                res = conn.execute(ins_sql, {
                    "id": str(uuid.uuid4()),
                    "nar_id": nar_id,
                    "application_name": application_name,
                    "release_number": release_number,
                    "rtype": rtype_text,
                    "doc_version": None,
                    "doc_hash": document_hash,
                    "doc_id": doc_id,
                    "doc_uri": doc_uri,
                    "page_num": source.page_num,
                    "embedding": _vec_literal(emb),
                    "chunk_index": chunk_index,
                    "chunk": record.text,
                    "chunk_hash": chunk_hash,
                    "metadata": json.dumps(md),
                })
                if res.first() is not None:
                    inserted += 1

    if ext == ".pdf":
        page_sources = pdf_pages_text(file_path)
        if not page_sources:
            logging.warning(f"[TEXT] No extractable PDF text found in: {file_path}")
        if pdf_page_aware:
            for source in page_sources:
                source_hash = persist_source_text(
                    nar_id,
                    release_number,
                    doc_id,
                    doc_uri,
                    document_hash,
                    source,
                    _build_source_metadata(source),
                )
                chunk_records = build_chunk_records(
                    source.text,
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                    merge_threshold=merge_threshold,
                    semantic_merge=semantic_merge,
                )
                _insert_chunk_records(source, source_hash, chunk_records)
            logging.info(f"[TEXT] Inserted {inserted}/{chunk_attempts} chunks from: {file_path}")
            return inserted
        source = build_pdf_document_source(page_sources)
        if source is None:
            logging.warning(f"[TEXT] No extractable PDF text found in: {file_path}")
            logging.info(f"[TEXT] Inserted {inserted}/{chunk_attempts} chunks from: {file_path}")
            return inserted
        source_hash = persist_source_text(
            nar_id,
            release_number,
            doc_id,
            doc_uri,
            document_hash,
            source,
            _build_source_metadata(source, {"page_count": len(page_sources)}),
        )
        chunk_records = build_chunk_records(
            source.text,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            merge_threshold=merge_threshold,
            semantic_merge=semantic_merge,
        )
        _insert_chunk_records(source, source_hash, chunk_records)
        logging.info(f"[TEXT] Inserted {inserted}/{chunk_attempts} chunks from: {file_path}")
        return inserted
    elif ext == ".docx":
        source = ExtractedTextSource(
            text_scope="document",
            page_num=None,
            text=docx_text(file_path),
            extractor_name="python-docx",
            parser_version=DOCX_VERSION,
        )
    elif ext == ".txt":
        source = ExtractedTextSource(
            text_scope="document",
            page_num=None,
            text=txt_text(file_path),
            extractor_name="plain-text",
            parser_version=PYTHON_RUNTIME_VERSION,
        )
    elif ext == ".csv":
        csv_chunks = split_csv_rows(file_path, rows_per_chunk=csv_rows_per_chunk)
        for csv_chunk in csv_chunks:
            source = ExtractedTextSource(
                text_scope="csv-chunk",
                page_num=None,
                text=csv_chunk.text,
                extractor_name="python-csv",
                parser_version=PYTHON_RUNTIME_VERSION,
            )
            extra_md = {
                "row_start": csv_chunk.row_start,
                "row_end": csv_chunk.row_end,
            }
            source_hash = persist_source_text(
                nar_id,
                release_number,
                doc_id,
                doc_uri,
                document_hash,
                source,
                _build_source_metadata(source, extra_md),
            )
            _insert_chunk_records(
                source,
                source_hash,
                [ChunkRecord(text=csv_chunk.text, start_offset=0, end_offset=len(csv_chunk.text))],
                chunk_index_base=csv_chunk.chunk_index,
                extra_md=extra_md,
            )
        logging.info(f"[TEXT] Inserted {inserted}/{chunk_attempts} chunks from: {file_path}")
        return inserted
    else:
        return inserted

    if source.text and source.text.strip():
        source_hash = persist_source_text(
            nar_id,
            release_number,
            doc_id,
            doc_uri,
            document_hash,
            source,
            _build_source_metadata(source),
        )
        chunk_records = build_chunk_records(
            source.text,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            merge_threshold=merge_threshold,
            semantic_merge=semantic_merge,
        )
        _insert_chunk_records(source, source_hash, chunk_records)
    else:
        logging.warning(f"[TEXT] No extractable text found in: {file_path}")

    logging.info(f"[TEXT] Inserted {inserted}/{chunk_attempts} chunks from: {file_path}")
    
    return inserted


def ingest_evidence_images_file(
    root: str,
    file_path: str,
    nar_id: str,
    release_number: str,
    source_id: str,
    category: Optional[str],
    rel_path: str,
    doc_id: str,
    doc_uri: str,
    rtype_image: str,
    mm_dim: int = MM_DIM,
    csv_image_columns: Optional[List[str]] = None,
    csv_caption_column: Optional[str] = None,
) -> int:
    """Inserts real evidence images into vs_mm_evidence_assets (NO snapshots).
    
    Logs:
      - how many images were found per file
      - each image row inserted (caption/page_num/chunk_index)
    """
    engine = get_engine()
    inserted = 0
    document_hash = doc_id
    ins_sql = text(f"""
        INSERT INTO {TABLE_MM_EVIDENCE}
        (id, nar_id, release_number, rtype, doc_version, doc_hash, doc_id, doc_uri, 
         caption, page_num, embedding, chunk_index, chunk, chunk_hash, metadata)
        VALUES (:id, :nar_id, :release_number, :rtype, :doc_version, :doc_hash, 
                :doc_id, :doc_uri, :caption, :page_num, CAST(:embedding AS vector), 
                :chunk_index, :chunk, :chunk_hash, CAST(:metadata AS jsonb))
        ON CONFLICT (chunk_hash) DO NOTHING
        RETURNING chunk_hash
    """)
    
    ext = os.path.splitext(file_path)[1].lower()
    image_extractor_name, image_parser_version = _image_extractor_details(ext)
    common_md = {
        "kind": "evidence-image",
        "rtype": rtype_image,
        "nar_id": nar_id,
        "release_number": release_number,
        "source_id": source_id,
        "category": category,
        "doc_id": doc_id,
        "doc_hash": document_hash,
        "doc_uri": doc_uri,
        "file_name": os.path.basename(file_path),
        "file_ext": ext.lstrip("."),
        "rel_path": rel_path,
        "extractor_name": image_extractor_name,
        "parser_version": image_parser_version,
    }
    
    images: List[ExtractedImageAsset] = []
    
    if ext == ".pdf":
        images.extend(extract_pdf_images(file_path))
    elif ext == ".docx":
        for chunk_index, img_bytes, caption in extract_docx_images(file_path):
            images.append(
                ExtractedImageAsset(
                    page_num=None,
                    chunk_index=chunk_index,
                    img_bytes=img_bytes,
                    caption=caption,
                    asset_label=caption,
                )
            )
    elif ext == ".csv":
        csv_imgs = parse_csv_image_rows(
            file_path,
            image_columns=csv_image_columns or [],
            caption_column=csv_caption_column,
        )
        for row_num, idx_in_row, img_bytes, caption in csv_imgs:
            asset_label = f"row{row_num}_{caption}"
            images.append(
                ExtractedImageAsset(
                    page_num=None,
                    chunk_index=idx_in_row,
                    img_bytes=img_bytes,
                    caption=asset_label,
                    asset_label=asset_label,
                )
            )
    else:
        return 0
    
    if not images:
        logging.debug(f"[IMG] No embedded images found in: {file_path}")
        return 0
    
    logging.info(f"[IMG] {len(images)} embedded images found in: {file_path}")
    
    with engine.begin() as conn:
        for asset in images:
            vec = mm_embed_image_with_caption(asset.img_bytes, caption=asset.caption, dim=mm_dim)
            md = dict(common_md)
            md.update({
                "page_num": asset.page_num,
                "chunk_index": asset.chunk_index,
                "caption": asset.caption,
                "asset_label": asset.asset_label,
                "nearby_heading": asset.nearby_heading,
                "preceding_text": asset.preceding_text,
                "page_text_excerpt": asset.page_text_excerpt,
            })
            chunk_hash = compute_image_chunk_hash(
                nar_id,
                release_number,
                document_hash,
                asset.page_num,
                asset.chunk_index,
                asset.caption,
                asset.img_bytes,
            )
            res = conn.execute(ins_sql, {
                "id": str(uuid.uuid4()),
                "nar_id": nar_id,
                "release_number": release_number,
                "rtype": rtype_image,
                "doc_version": None,
                "doc_hash": document_hash,
                "doc_id": doc_id,
                "doc_uri": doc_uri,
                "caption": asset.caption,
                "page_num": asset.page_num,
                "embedding": _vec_literal(vec),
                "chunk_index": asset.chunk_index,
                "chunk": None,
                "chunk_hash": chunk_hash,
                "metadata": json.dumps(md),
            })
            if res.first() is not None:
                inserted += 1
    
    logging.info(f"[IMG] Inserted {inserted}/{len(images)} images from: {file_path}")
    return inserted


# ============================================================================
# Purge Helpers
# ============================================================================

def purge_by_scope(nar_id: str, release_number: str):
    """Purges evidence text+images for a given scope."""
    engine = get_engine()
    with engine.begin() as conn:
        params = {"nar": nar_id, "rel": release_number}
        t1 = conn.execute(
            text(f"DELETE FROM {TABLE_EVIDENCE} WHERE nar_id=:nar AND release_number=:rel"),
            params,
        ).rowcount
        t2 = conn.execute(
            text(f"DELETE FROM {TABLE_MM_EVIDENCE} WHERE nar_id=:nar AND release_number=:rel"),
            params,
        ).rowcount
        t3 = conn.execute(
            text(f"DELETE FROM {TABLE_EVIDENCE_TEXT_SOURCES} WHERE nar_id=:nar AND release_number=:rel"),
            params,
        ).rowcount
    logging.info(
        f"[PURGE] nar_id={nar_id}, release_number={release_number} "
        f"deleted text_rows={t1 or 0}, image_rows={t2 or 0}, source_rows={t3 or 0}"
    )


def purge_all_tables():
    """Purges ALL rows from BOTH evidence tables:
    
      - vs_evidence
      - vs_mm_evidence_assets
    
    Use with caution.
    """
    engine = get_engine()
    with engine.begin() as conn:
        t1 = conn.execute(text(f"DELETE FROM {TABLE_EVIDENCE}")).rowcount
        t2 = conn.execute(text(f"DELETE FROM {TABLE_MM_EVIDENCE}")).rowcount
        t3 = conn.execute(text(f"DELETE FROM {TABLE_EVIDENCE_TEXT_SOURCES}")).rowcount
    logging.warning(
        f"[PURGE-ALL] Deleted ALL rows: "
        f"vs_evidence={t1 or 0}, vs_mm_evidence_assets={t2 or 0}, "
        f"vs_evidence_text_sources={t3 or 0}"
    )


# ============================================================================
# CLI
# ============================================================================

def main():
    """Command-line interface."""
    ap = argparse.ArgumentParser(
        description="Evidence ingestion (text+images) into pgvector tables"
    )
    
    # Core arguments for ingestion
    ap.add_argument(
        "--root",
        help="Evidence root folder. Must contain subfolders as source buckets.",
    )
    ap.add_argument("--nar-id", help="NAR/Change identifier for the evidence.")
    ap.add_argument("--application-name", help="Application name for the evidence.")
    ap.add_argument("--release-number", help="Release number for the evidence.")
    
    # Options
    ap.add_argument(
        "--uri-prefix",
        default="",
        help="Optional prefix for doc_uri (e.g., gs://bucket/evidence)",
    )
    ap.add_argument(
        "--enable-mm",
        action="store_true",
        help="Also ingest evidence images/diagrams",
    )
    ap.add_argument(
        "--mm-dim",
        type=int,
        default=MM_DIM,
        help="MM embedding dim (must match VECTOR(1408))",
    )
    ap.add_argument("--chunk-size", type=int, default=1000)
    ap.add_argument("--chunk-overlap", type=int, default=100)
    ap.add_argument(
        "--semantic-merge",
        action="store_true",
        help="Merge adjacent deterministic chunks using embedding similarity.",
    )
    ap.add_argument(
        "--merge-threshold",
        type=float,
        default=0.85,
        help="Similarity threshold used only with --semantic-merge.",
    )
    ap.add_argument(
        "--no-pdf-page-aware",
        action="store_true",
        help="If set, PDF text ingests as whole doc (page_num NULL)",
    )
    ap.add_argument(
        "--rtype-text",
        default="new",
        help="rtype for text. Allowed: indev, new, migration, rehost, normal",
    )
    ap.add_argument(
        "--rtype-image",
        default="new",
        help="rtype for images. Allowed: indev, new, migration, rehost, normal",
    )
    ap.add_argument(
        "--strict-root",
        action="store_true",
        help="Fail if any file is directly under root (no source folder)",
    )
    ap.add_argument(
        "--csv-rows-per-chunk",
        type=int,
        default=50,
        help="Rows per chunk when ingesting CSV files (header repeated).",
    )
    ap.add_argument(
        "--csv-image-column",
        action="append",
        default=[],
        help="CSV column name(s) containing base64-encoded images; repeatable.",
    )
    ap.add_argument(
        "--csv-caption-column",
        default="",
        help="Optional CSV column name to use as caption for images.",
    )
    
    # Purge-only arguments
    ap.add_argument(
        "--purge-scope",
        action="store_true",
        help="Purge all evidence for the specified --nar-id and --release-number and EXIT.",
    )
    ap.add_argument(
        "--purge-all",
        action="store_true",
        help="Delete ALL rows from BOTH evidence tables and EXIT (no ingestion)",
    )
    
    # Filtering arguments
    ap.add_argument(
        "--only-source-id",
        action="append",
        default=[],
        help="Ingest only files whose derived source_id matches this value. Repeatable.",
    )
    ap.add_argument(
        "--only-file",
        action="append",
        default=[],
        help="Ingest only these files (absolute or relative to root). Repeatable.",
    )
    
    args = ap.parse_args()
    
    rtype_text = resolve_rtype(args.rtype_text)
    rtype_image = resolve_rtype(args.rtype_image)
    logging.info(f"[CFG] Using rtype-text='{rtype_text}', rtype-image='{rtype_image}'")
    
    # Fail fast
    get_engine()
    ensure_support_tables()
    
    if EMBED_DIM is None:
        logging.warning(
            f"Unknown embedding model '{EMBED_MODEL}'. "
            f"Ensure VECTOR(dim) matches model output."
        )
    
    if args.mm_dim != MM_DIM:
        logging.info(
            f"[INFO] Using MM-DIM override {args.mm_dim} "
            f"(DDL expects 1408 by default)."
        )
    if args.semantic_merge:
        logging.warning(
            "[CFG] Semantic merge is enabled. Chunk boundaries may change with embedding behavior, "
            "but stored chunk text and offsets remain exact source slices."
        )
    
    # PURGE ONLY
    if args.purge_all:
        logging.warning(
            "[PURGE-ALL] Purging ALL rows from evidence tables and exiting (no ingestion)."
        )
        purge_all_tables()
        return
    
    if args.purge_scope:
        if not (args.nar_id and args.release_number):
            raise SystemExit("ERROR: --nar-id and --release-number are required for --purge-scope.")
        logging.info(
            f"[PURGE] Purging evidence for nar_id={args.nar_id}, "
            f"release_number={args.release_number} and exiting."
        )
        purge_by_scope(args.nar_id, args.release_number)
        return
    
    # From here on, we expect ingestion flags; root becomes required
    if not all([args.root, args.nar_id, args.application_name, args.release_number]):
        raise SystemExit(
            "ERROR: --root, --nar-id, --application-name, and --release-number "
            "are required for ingestion."
        )

    validate_runtime_contract(args.mm_dim)
    
    uri_prefix = args.uri_prefix.strip() or None
    pdf_page_aware = not args.no_pdf_page_aware
    
    # Normalize source-id filter (case-insensitive compare on derived source_id)
    allowed_source_ids = {s.strip().lower() for s in (args.only_source_id or []) if s.strip()}
    
    # If explicit file(s) provided, use only those; else walk root
    if args.only_file:
        files = _resolve_file_paths(args.root, args.only_file)
        logging.info(f"[SCAN] Using explicit files: {len(args.only_file)} provided, {len(files)} valid.")
    else:
        files = list(iter_guidance_files(args.root))
        logging.info(f"[SCAN] Discovered {len(files)} evidence files under {args.root}")
    
    # Optional filter by source_id
    if allowed_source_ids:
        filtered_files: List[str] = []
        for fp in files:
            try:
                source_id, _, _ = derive_source_and_category(
                    args.root, fp, strict_root=args.strict_root
                )
            except Exception as e:
                logging.warning(f"[FILTER] Skipping {fp}: {e}")
                continue
            if source_id != "UNKNOWN" and source_id.strip().lower() in allowed_source_ids:
                filtered_files.append(fp)
        logging.info(f"[FILTER] source_id filter kept {len(filtered_files)}/{len(files)} files.")
        files = filtered_files
    
    if not files:
        raise SystemExit(
            "Nothing to ingest after applying filters. "
            "Check --root, --only-file, and/or --only-source-id."
        )
    
    total_text = 0
    total_images = 0
    
    for fp in files:
        # Derive source/category; enforce structure if requested
        try:
            source_id, category, rel_path = derive_source_and_category(
                args.root, fp, strict_root=args.strict_root
            )
        except Exception as e:
            logging.error(f"[ERROR] Skipping {fp}: {e}")
            continue
        
        if source_id == "UNKNOWN":
            continue
        
        doc_id = compute_doc_id(fp)
        
        doc_uri = compute_doc_uri(args.root, fp, uri_prefix)
        logging.info(
            f"[EVIDENCE] file={rel_path} source_id={source_id} category={category} "
            f"doc_id={doc_id[:12]}..."
        )
        
        inst_t = ingest_evidence_text_file(
            root=args.root,
            file_path=fp,
            nar_id=args.nar_id,
            application_name=args.application_name,
            release_number=args.release_number,
            source_id=source_id,
            category=category,
            rel_path=rel_path,
            doc_id=doc_id,
            doc_uri=doc_uri,
            rtype_text=rtype_text,
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap,
            merge_threshold=args.merge_threshold,
            semantic_merge=args.semantic_merge,
            pdf_page_aware=pdf_page_aware,
            csv_rows_per_chunk=args.csv_rows_per_chunk,
        )
        total_text += inst_t
        
        if args.enable_mm:
            inst_i = ingest_evidence_images_file(
                root=args.root,
                file_path=fp,
                nar_id=args.nar_id,
                release_number=args.release_number,
                source_id=source_id,
                category=category,
                rel_path=rel_path,
                doc_id=doc_id,
                doc_uri=doc_uri,
                rtype_image=rtype_image,
                mm_dim=args.mm_dim,
                csv_image_columns=[c for c in args.csv_image_column if c.strip()],
                csv_caption_column=(args.csv_caption_column or "").strip() or None,
            )
            total_images += inst_i
    
    logging.info(f"[DONE] inserted_text_chunks={total_text} inserted_evidence_images={total_images}")


if __name__ == "__main__":
    main()
